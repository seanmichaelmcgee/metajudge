#!/usr/bin/env python3
"""Generate three QA audit .docx documents for MetaJudge v0.5.5.1.

Doc 1: Complete audit (all 201 items)
Doc 2: Highest-yield items (best discrimination across models)
Doc 3: Most suspect/problematic items

Usage:
    python scripts/generate_qa_audit_docx.py
"""

import csv
import json
import statistics
import sys
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Paths ──────────────────────────────────────────────────────────────────
RESULTS_DIR = Path("/tmp/v0551/v0.5.5.1 - results")
KAGGLE_DIR = Path("/tmp/v0551-kaggle/v0.5.5.1 - kaggle-upload")
OUTPUT_DIR = Path("/home/user/metajudge/outputs")

CAL_CSV = RESULTS_DIR / "calibration_item_audit (8).csv"
FB_CSV = RESULTS_DIR / "family_b_item_audit (8).csv"
RUN_SUMMARY = RESULTS_DIR / "run_summary (7).json"
BRIDGE_ALL = RESULTS_DIR / "bridge_report_all_models (2).json"
SUCCESS_CRIT = RESULTS_DIR / "success_criteria_verdict.json"
REVIEW_QUEUE = RESULTS_DIR / "audit_review_queue (2).csv"
CAL_BY_MECH = RESULTS_DIR / "calibration_summary_by_mechanism (2).csv"
FB_BY_MODEL = RESULTS_DIR / "family_b_summary_by_model (2).csv"

BENCHMARK_JSON = KAGGLE_DIR / "metajudge_benchmark_v1.json"
FAMILY_B_JSON = KAGGLE_DIR / "family_b_pilot_v2.json"
JUSTIFICATIONS_MD = OUTPUT_DIR / "metajudge_v0551_gold_answer_justifications.md"

# ── Display names ──────────────────────────────────────────────────────────
MODEL_ORDER = [
    "google/gemini-2.5-flash",
    "google/gemini-2.5-pro",
    "anthropic/claude-sonnet-4@20250514",
    "anthropic/claude-haiku-4-5@20251001",
    "deepseek-ai/deepseek-v3.1",
]
MODEL_SHORT = {
    "google/gemini-2.5-flash": "Gemini Flash",
    "google/gemini-2.5-pro": "Gemini Pro",
    "anthropic/claude-sonnet-4@20250514": "Sonnet 4",
    "anthropic/claude-haiku-4-5@20251001": "Haiku 4.5",
    "deepseek-ai/deepseek-v3.1": "DeepSeek V3.1",
}

# ── Colors ─────────────────────────────────────────────────────────────────
GREEN_BG = "C6EFCE"
RED_BG = "FFC7CE"
ORANGE_BG = "FFE0B2"
HEADER_BG = "1F4E79"
HEADER_FG = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY_BG = "F2F2F2"
GOLD_ANS_BG = "FFF9C4"
JUSTIFICATION_BG = "E8F5E9"


# ── Data loading ───────────────────────────────────────────────────────────

def load_justifications():
    """Parse justifications markdown into dict keyed by item_id."""
    if not JUSTIFICATIONS_MD.exists():
        print("  Warning: justifications file not found, skipping.")
        return {}
    import re
    with open(JUSTIFICATIONS_MD) as f:
        content = f.read()
    # Match #### item_id blocks and extract the Justification line
    pattern = r'#### (\S+)\n.*?\*\*Justification:\*\* (.+?)(?=\n\n#### |\n\n### |\n\n## |\n\n---|\Z)'
    matches = re.findall(pattern, content, re.DOTALL)
    result = {}
    for item_id, justification in matches:
        result[item_id] = justification.strip()
    return result


def load_benchmark_items():
    """Load full question text and metadata from benchmark JSONs."""
    cal_items = {}
    with open(BENCHMARK_JSON) as f:
        for item in json.load(f):
            cal_items[item["item_id"]] = item

    fb_items = {}
    with open(FAMILY_B_JSON) as f:
        for item in json.load(f):
            fb_items[item["item_id"]] = item

    return cal_items, fb_items


def load_csv(path):
    """Load CSV with DictReader."""
    with open(path) as f:
        return list(csv.DictReader(f))


def load_json(path):
    with open(path) as f:
        return json.load(f)


def build_cal_grouped(cal_rows, cal_meta):
    """Group calibration results by item_id, enriched with metadata."""
    items = defaultdict(dict)
    for row in cal_rows:
        iid = row["item_id"]
        model = row["model_name"]
        items[iid][model] = row
    # Build item list with metadata
    result = []
    for iid, models_data in sorted(items.items()):
        meta = cal_meta.get(iid, {})
        result.append({
            "item_id": iid,
            "family": "A",
            "family_name": "Confidence Calibration",
            "question": meta.get("question", next(iter(models_data.values())).get("question", "")),
            "gold_answer": meta.get("gold_answer", next(iter(models_data.values())).get("gold_answer", "")),
            "aliases": meta.get("aliases", []),
            "mechanism": meta.get("mechanism_primary", next(iter(models_data.values())).get("mechanism", "")),
            "cue_type": meta.get("cue_type", ""),
            "grader_rule": meta.get("rule", next(iter(models_data.values())).get("grader_rule", "")),
            "stratum": meta.get("stratum", ""),
            "final_classification": meta.get("final_classification", ""),
            "temporal_brittle": meta.get("temporal_brittle", False),
            "controversy_risk": meta.get("controversy_risk", False),
            "difficulty": "",
            "category": "",
            "gold_action": "",
            "is_false_presupposition": False,
            "models": models_data,
        })
    return result


def build_fb_grouped(fb_rows, fb_meta):
    """Group Family B results by item_id, enriched with metadata."""
    items = defaultdict(dict)
    for row in fb_rows:
        iid = row["item_id"]
        model = row["model_name"]
        items[iid][model] = row
    result = []
    for iid, models_data in sorted(items.items()):
        meta = fb_meta.get(iid, {})
        result.append({
            "item_id": iid,
            "family": "B",
            "family_name": "Selective Abstention",
            "question": meta.get("question", next(iter(models_data.values())).get("question", "")),
            "gold_answer": meta.get("gold_answer", next(iter(models_data.values())).get("gold_answer", "")),
            "aliases": [],
            "mechanism": "",
            "cue_type": "",
            "grader_rule": "",
            "stratum": "",
            "final_classification": "",
            "temporal_brittle": False,
            "controversy_risk": False,
            "difficulty": meta.get("difficulty", ""),
            "category": meta.get("category", ""),
            "gold_action": meta.get("gold_action", next(iter(models_data.values())).get("gold_action", "")),
            "is_false_presupposition": meta.get("is_false_presupposition", False),
            "acceptable_actions": meta.get("acceptable_actions", []),
            "premise_type": meta.get("premise_type", ""),
            "models": models_data,
        })
    return result


# ── Selection logic ────────────────────────────────────────────────────────

def compute_discrimination(item):
    """Compute discrimination score for an item (higher = more discriminating)."""
    if item["family"] == "A":
        vals = []
        for m in MODEL_ORDER:
            row = item["models"].get(m)
            if row:
                vals.append(float(row.get("brier_score", 0)))
        if len(vals) >= 2:
            return statistics.stdev(vals)
    else:
        vals = []
        for m in MODEL_ORDER:
            row = item["models"].get(m)
            if row:
                vals.append(float(row.get("utility", 0)))
        if len(vals) >= 2:
            return statistics.stdev(vals)
    return 0.0


def compute_agreement_rate(item):
    """Fraction of models that got the item correct."""
    correct = 0
    total = 0
    for m in MODEL_ORDER:
        row = item["models"].get(m)
        if row:
            total += 1
            if row.get("is_correct", "").lower() == "true":
                correct += 1
    return correct / total if total > 0 else 0.5


def select_highest_yield(all_items, top_n=40):
    """Select items that best discriminate between models."""
    scored = [(compute_discrimination(item), item) for item in all_items]
    scored.sort(key=lambda x: -x[0])
    return [item for _, item in scored[:top_n]]


def select_suspect(all_items, review_queue_rows):
    """Select items that are suspect or problematic."""
    # Build set of flagged item_ids from review queue
    review_flagged = set()
    review_details = {}
    for row in review_queue_rows:
        iid = row.get("item_id", "")
        review_flagged.add(iid)
        detail = review_details.get(iid, [])
        detail.append(f"{row.get('priority', '')} ({row.get('model', '')}): {row.get('detail', '')}")
        review_details[iid] = detail

    selected = []
    for item in all_items:
        reasons = []
        iid = item["item_id"]

        # 1. In review queue
        if iid in review_flagged:
            reasons.append("Flagged in review queue: " + "; ".join(review_details.get(iid, [])[:2]))

        # 2. Majority of models disagree with gold
        agreement = compute_agreement_rate(item)
        if agreement < 0.4:
            n_wrong = sum(1 for m in MODEL_ORDER if item["models"].get(m, {}).get("is_correct", "").lower() == "false")
            reasons.append(f"Majority wrong ({n_wrong} models disagree with gold)")

        # 3. Temporal brittle or controversy risk
        if item.get("temporal_brittle"):
            reasons.append("Temporal brittle")
        if item.get("controversy_risk"):
            reasons.append("Controversy risk")

        # 4. False presupposition (Family B)
        if item.get("is_false_presupposition"):
            reasons.append("False presupposition")

        # 5. Borderline classification
        if item.get("final_classification") == "BORDERLINE":
            reasons.append("Borderline classification")

        # 6. Best model (Gemini Pro) wrong
        pro_row = item["models"].get("google/gemini-2.5-pro")
        if pro_row and pro_row.get("is_correct", "").lower() == "false":
            reasons.append("Best model (Gemini Pro) incorrect")

        # 7. Ambiguity or parser risk
        for m in MODEL_ORDER:
            row = item["models"].get(m, {})
            if row.get("ambiguity_risk", "").lower() == "true":
                reasons.append("Ambiguity risk flagged")
                break
            if row.get("parser_risk", "").lower() == "true":
                reasons.append("Parser risk flagged")
                break

        if reasons:
            item["_suspect_reasons"] = reasons
            selected.append(item)

    return selected


# ── Document generation helpers ────────────────────────────────────────────

def set_cell_bg(cell, color_hex):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, size=8, color=None, align=None):
    """Set text in a cell with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if bold:
        run.font.bold = True
    if color:
        run.font.color.rgb = color
    # Reduce paragraph spacing
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after = Pt(1)


def add_toc(doc):
    """Add a Table of Contents field that Word will populate on open."""
    p = doc.add_paragraph()
    run = p.add_run()
    fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fld_char_begin)

    run2 = p.add_run()
    instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    run2._r.append(instr)

    run3 = p.add_run()
    fld_char_separate = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._r.append(fld_char_separate)

    run4 = p.add_run("(Right-click and select 'Update Field' to populate Table of Contents)")
    run4.font.size = Pt(10)
    run4.font.italic = True
    run4.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    run5 = p.add_run()
    fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._r.append(fld_char_end)


def add_cover_page(doc, title, subtitle, run_summary, success_criteria):
    """Add a cover page with title and key metadata."""
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run.font.name = "Calibri"

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    run.font.name = "Calibri"

    doc.add_paragraph()  # spacer

    # Metadata
    meta_lines = [
        f"Benchmark Version: v0.5.5.1 (grading engine: {run_summary.get('grading_engine', 'grading_v2')})",
        f"Date: 2026-03-25",
        f"Calibration Items: {run_summary.get('calibration', {}).get('total_items', 117)}",
        f"Family B Items: {run_summary.get('family_b', {}).get('total_items', 84)}",
        f"Models: {run_summary.get('calibration', {}).get('models', 5)}",
        f"Success Criteria Verdict: {success_criteria.get('verdict', 'N/A')} ({success_criteria.get('total_pass', 0)}/{success_criteria.get('total_criteria', 5)} passing)",
    ]
    for line in meta_lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = "Calibri"

    doc.add_page_break()


def add_executive_summary(doc, run_summary, success_criteria):
    """Add executive summary with leaderboard and success criteria."""
    doc.add_heading("Executive Summary", level=1)

    # Calibration leaderboard
    doc.add_heading("Calibration Leaderboard", level=2)
    cal_models = run_summary.get("calibration", {}).get("per_model", {})

    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Model", "Items", "Accuracy", "1-Brier", "ECE", "Conf. Disc.", "Overconf. Rate"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=9, color=HEADER_FG)
        set_cell_bg(table.rows[0].cells[i], HEADER_BG)

    for model_key in MODEL_ORDER:
        mdata = cal_models.get(model_key, {})
        if not mdata:
            continue
        row = table.add_row()
        vals = [
            MODEL_SHORT.get(model_key, model_key),
            str(mdata.get("n_items", "")),
            f"{mdata.get('accuracy', 0):.1%}",
            f"{mdata.get('mean_1_minus_brier', 0):.4f}",
            f"{mdata.get('ece', 0):.4f}",
            f"{mdata.get('confidence_discrimination', 0):.4f}",
            f"{mdata.get('overconfidence_rate', 0):.1%}",
        ]
        for i, v in enumerate(vals):
            set_cell_text(row.cells[i], v, size=9)

    doc.add_paragraph()

    # Family B leaderboard
    doc.add_heading("Family B Leaderboard", level=2)
    fb_models = run_summary.get("family_b", {}).get("per_model", {})

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Model", "Items", "Action Acc.", "Mean Utility", "Over-Answer", "Verify Trigger"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=9, color=HEADER_FG)
        set_cell_bg(table.rows[0].cells[i], HEADER_BG)

    for model_key in MODEL_ORDER:
        mdata = fb_models.get(model_key, {})
        if not mdata or mdata.get("n_items", 0) < 2:
            continue
        row = table.add_row()
        vals = [
            MODEL_SHORT.get(model_key, model_key),
            str(mdata.get("n_items", "")),
            f"{mdata.get('action_accuracy', 0):.1%}",
            f"{mdata.get('mean_utility', 0):.4f}",
            f"{mdata.get('over_answer_rate', 0):.1%}",
            f"{mdata.get('verify_trigger_rate', 0):.1%}",
        ]
        for i, v in enumerate(vals):
            set_cell_text(row.cells[i], v, size=9)

    doc.add_paragraph()

    # Success criteria
    doc.add_heading("Success Criteria", level=2)
    criteria = success_criteria.get("criteria", {})
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Criterion", "Name", "Threshold", "Measured", "Pass"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=9, color=HEADER_FG)
        set_cell_bg(table.rows[0].cells[i], HEADER_BG)

    for cid in ["C1", "C2", "C3", "C4", "C5"]:
        c = criteria.get(cid, {})
        row = table.add_row()
        passed = str(c.get("pass", "")).lower() == "true"
        vals = [
            cid,
            c.get("name", ""),
            str(c.get("threshold", "")),
            str(c.get("measured", "")),
            "PASS" if passed else "FAIL",
        ]
        for i, v in enumerate(vals):
            set_cell_text(row.cells[i], v, size=9)
        set_cell_bg(row.cells[4], GREEN_BG if passed else RED_BG)

    doc.add_page_break()


def truncate_text(text, max_len=80):
    """Truncate text for headings."""
    text = str(text).replace("\n", " ").strip()
    if len(text) > max_len:
        return text[:max_len] + "..."
    return text


def add_item_section(doc, item, heading_level=3, show_suspect_reasons=False):
    """Add a single item section to the document."""
    iid = item["item_id"]
    q_short = truncate_text(item["question"])

    # Heading
    doc.add_heading(f"{iid} -- {q_short}", level=heading_level)

    # Metadata line
    meta_parts = [f"Family {item['family']}: {item['family_name']}"]
    if item.get("mechanism"):
        meta_parts.append(f"Mechanism: {item['mechanism']}")
    if item.get("cue_type"):
        meta_parts.append(f"Cue: {item['cue_type']}")
    if item.get("category"):
        meta_parts.append(f"Category: {item['category']}")
    if item.get("grader_rule"):
        meta_parts.append(f"Grader: {item['grader_rule']}")
    if item.get("difficulty"):
        meta_parts.append(f"Difficulty: {item['difficulty']}")
    if item.get("gold_action"):
        meta_parts.append(f"Gold Action: {item['gold_action']}")
    if item.get("final_classification"):
        meta_parts.append(f"Classification: {item['final_classification']}")
    if item.get("premise_type"):
        meta_parts.append(f"Premise: {item['premise_type']}")

    flags = []
    if item.get("temporal_brittle"):
        flags.append("TEMPORAL-BRITTLE")
    if item.get("controversy_risk"):
        flags.append("CONTROVERSY-RISK")
    if item.get("is_false_presupposition"):
        flags.append("FALSE-PRESUPPOSITION")
    if flags:
        meta_parts.append(f"Flags: {', '.join(flags)}")

    p = doc.add_paragraph(" | ".join(meta_parts))
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        run.font.name = "Calibri"

    # Suspect reasons (for Doc 3)
    if show_suspect_reasons and item.get("_suspect_reasons"):
        p = doc.add_paragraph()
        run = p.add_run("Concerns: ")
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        run.font.name = "Calibri"
        for reason in item["_suspect_reasons"]:
            run = p.add_run(f"\n  - {reason}")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            run.font.name = "Calibri"

    # Question text (shaded box)
    p = doc.add_paragraph()
    run = p.add_run("Question: ")
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    run = p.add_run(str(item["question"]))
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    # Apply light background shading to the paragraph
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{LIGHT_GRAY_BG}" w:val="clear"/>')
    pPr.append(shading)

    # Gold answer
    p = doc.add_paragraph()
    run = p.add_run("Gold Answer: ")
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    run = p.add_run(str(item["gold_answer"]))
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x00, 0x66, 0x00)
    if item.get("aliases"):
        run = p.add_run(f"  (aliases: {', '.join(item['aliases'])})")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        run.font.name = "Calibri"
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{GOLD_ANS_BG}" w:val="clear"/>')
    pPr.append(shading)

    # Justification (if available)
    justification = item.get("_justification", "")
    if justification:
        p = doc.add_paragraph()
        run = p.add_run("Justification: ")
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
        run = p.add_run(justification)
        run.font.size = Pt(9)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
        pPr = p._p.get_or_add_pPr()
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{JUSTIFICATION_BG}" w:val="clear"/>')
        pPr.append(shading)

    # Model response table
    is_family_b = item["family"] == "B"

    # Determine which models have data
    available_models = [m for m in MODEL_ORDER if m in item["models"]]
    n_cols = len(available_models) + 1  # +1 for row label

    if is_family_b:
        row_labels = ["Decision", "Answer", "Confidence", "Correct?", "Utility"]
    else:
        row_labels = ["Answer", "Confidence", "Correct?", "Brier Score (1-Brier)"]

    table = doc.add_table(rows=len(row_labels) + 1, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    set_cell_text(table.rows[0].cells[0], "", bold=True, size=8)
    set_cell_bg(table.rows[0].cells[0], HEADER_BG)
    for ci, model in enumerate(available_models):
        set_cell_text(table.rows[0].cells[ci + 1], MODEL_SHORT.get(model, model), bold=True, size=8, color=HEADER_FG)
        set_cell_bg(table.rows[0].cells[ci + 1], HEADER_BG)

    # Data rows
    for ri, label in enumerate(row_labels):
        row_idx = ri + 1
        set_cell_text(table.rows[row_idx].cells[0], label, bold=True, size=8)
        set_cell_bg(table.rows[row_idx].cells[0], LIGHT_GRAY_BG)

        for ci, model in enumerate(available_models):
            mrow = item["models"].get(model, {})
            cell = table.rows[row_idx].cells[ci + 1]

            if label == "Answer":
                ans = str(mrow.get("model_answer", "N/A"))
                # Truncate very long answers for readability
                if len(ans) > 300:
                    ans = ans[:297] + "..."
                set_cell_text(cell, ans, size=8)

            elif label == "Decision":
                dec = str(mrow.get("model_decision", "N/A"))
                set_cell_text(cell, dec, size=8, bold=True)
                # Color code decisions
                if dec == item.get("gold_action", ""):
                    set_cell_bg(cell, GREEN_BG)
                elif dec in ("answer", "clarify", "verify", "abstain"):
                    set_cell_bg(cell, ORANGE_BG)

            elif label == "Confidence":
                conf = mrow.get("confidence", "N/A")
                try:
                    conf_f = float(conf)
                    set_cell_text(cell, f"{conf_f:.2f}", size=8)
                except (ValueError, TypeError):
                    set_cell_text(cell, str(conf), size=8)

            elif label == "Correct?":
                correct = str(mrow.get("is_correct", "")).lower() == "true"
                set_cell_text(cell, "Yes" if correct else "No", size=8, bold=True)
                set_cell_bg(cell, GREEN_BG if correct else RED_BG)

            elif label.startswith("Brier"):
                brier = mrow.get("brier_score", "N/A")
                try:
                    brier_f = float(brier)
                    set_cell_text(cell, f"{brier_f:.4f}", size=8)
                    if brier_f >= 0.9:
                        set_cell_bg(cell, GREEN_BG)
                    elif brier_f < 0.5:
                        set_cell_bg(cell, RED_BG)
                except (ValueError, TypeError):
                    set_cell_text(cell, str(brier), size=8)

            elif label == "Utility":
                util = mrow.get("utility", "N/A")
                try:
                    util_f = float(util)
                    set_cell_text(cell, f"{util_f:+.2f}", size=8)
                    if util_f >= 0.5:
                        set_cell_bg(cell, GREEN_BG)
                    elif util_f <= -0.5:
                        set_cell_bg(cell, RED_BG)
                    elif util_f < 0:
                        set_cell_bg(cell, ORANGE_BG)
                except (ValueError, TypeError):
                    set_cell_text(cell, str(util), size=8)

    # Scoring impact note
    correct_count = sum(
        1 for m in available_models
        if item["models"].get(m, {}).get("is_correct", "").lower() == "true"
    )
    total_models = len(available_models)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run(f"Cross-model agreement: {correct_count}/{total_models} correct")
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    run.font.name = "Calibri"

    if item["family"] == "A":
        scores = []
        for m in available_models:
            try:
                scores.append(float(item["models"][m].get("brier_score", 0)))
            except (ValueError, TypeError):
                pass
        if scores:
            spread = max(scores) - min(scores)
            run = p.add_run(f" | Brier spread: {spread:.4f}")
            run.font.size = Pt(8)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
            run.font.name = "Calibri"

    doc.add_paragraph()  # spacer


def group_items_by_domain(items):
    """Group items by family then mechanism/category for TOC structure."""
    groups = defaultdict(list)
    for item in items:
        if item["family"] == "A":
            key = ("A", item.get("mechanism", "Other"))
        else:
            cat = item.get("category", "Other")
            # Normalize very long category strings
            if len(cat) > 40:
                cat = cat[:37] + "..."
            key = ("B", cat)
        groups[key].append(item)
    return groups


def generate_document(items, title, subtitle, run_summary, success_criteria,
                      output_path, show_suspect_reasons=False):
    """Generate a complete .docx document."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    # Set narrow margins for more content per page
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Cover page
    add_cover_page(doc, title, subtitle, run_summary, success_criteria)

    # Table of Contents
    doc.add_heading("Table of Contents", level=1)
    add_toc(doc)
    doc.add_page_break()

    # Group items
    groups = group_items_by_domain(items)

    # Sort groups: Family A first, then Family B
    sorted_keys = sorted(groups.keys(), key=lambda k: (k[0], k[1]))

    current_family = None
    for (family, domain), group_items in sorted(groups.items(), key=lambda x: (x[0][0], x[0][1])):
        # Family heading
        if family != current_family:
            current_family = family
            family_name = "Family A: Confidence Calibration" if family == "A" else "Family B: Selective Abstention"
            doc.add_heading(family_name, level=1)

        # Domain heading
        doc.add_heading(f"{domain}", level=2)

        # Items
        for item in sorted(group_items, key=lambda x: x["item_id"]):
            add_item_section(doc, item, heading_level=3, show_suspect_reasons=show_suspect_reasons)

    # Save
    doc.save(str(output_path))
    print(f"  Saved: {output_path} ({len(items)} items)")


# ── Main ───────────────────────────────────────────────────────────────────

def attach_justifications(items, justifications):
    """Attach justification text to each item dict."""
    matched = 0
    for item in items:
        j = justifications.get(item["item_id"], "")
        if j:
            item["_justification"] = j
            matched += 1
    return matched


def main():
    import argparse
    parser = argparse.ArgumentParser(description="Generate QA audit .docx documents")
    parser.add_argument("--merged-only", action="store_true",
                        help="Only generate the merged (with justifications) document")
    args = parser.parse_args()

    print("Loading data...")
    cal_meta, fb_meta = load_benchmark_items()
    cal_rows = load_csv(CAL_CSV)
    fb_rows = load_csv(FB_CSV)
    run_summary = load_json(RUN_SUMMARY)
    success_criteria = load_json(SUCCESS_CRIT)
    review_queue = load_csv(REVIEW_QUEUE)

    print(f"  Calibration: {len(cal_rows)} rows, {len(cal_meta)} benchmark items")
    print(f"  Family B: {len(fb_rows)} rows, {len(fb_meta)} benchmark items")

    # Load justifications
    justifications = load_justifications()
    print(f"  Justifications: {len(justifications)} loaded")

    # Build grouped data
    cal_items = build_cal_grouped(cal_rows, cal_meta)
    fb_items = build_fb_grouped(fb_rows, fb_meta)
    all_items = cal_items + fb_items

    # Attach justifications
    matched = attach_justifications(all_items, justifications)
    print(f"  Justifications matched: {matched}/{len(all_items)}")

    print(f"  Total unique items: {len(all_items)} (Cal: {len(cal_items)}, FB: {len(fb_items)})")

    OUTPUT_DIR.mkdir(exist_ok=True)

    if not args.merged_only:
        # Doc 1: Complete audit (without justifications — original)
        # Strip justifications for the plain version
        for item in all_items:
            item.pop("_justification", None)

        print("\nGenerating Doc 1: Complete Audit...")
        generate_document(
            items=all_items,
            title="MetaJudge v0.5.5.1",
            subtitle="Complete QA Audit Report -- All Items",
            run_summary=run_summary,
            success_criteria=success_criteria,
            output_path=OUTPUT_DIR / "metajudge_v0551_qa_audit_complete.docx",
        )

        # Doc 2: Highest yield
        print("\nGenerating Doc 2: Highest Yield...")
        highest_yield = select_highest_yield(all_items, top_n=40)
        generate_document(
            items=highest_yield,
            title="MetaJudge v0.5.5.1",
            subtitle="Highest-Yield Items -- Best Model Discrimination",
            run_summary=run_summary,
            success_criteria=success_criteria,
            output_path=OUTPUT_DIR / "metajudge_v0551_qa_audit_highest_yield.docx",
        )

        # Doc 3: Most suspect
        print("\nGenerating Doc 3: Most Suspect...")
        suspect = select_suspect(all_items, review_queue)
        generate_document(
            items=suspect,
            title="MetaJudge v0.5.5.1",
            subtitle="Suspect & Problematic Items -- Review Required",
            run_summary=run_summary,
            success_criteria=success_criteria,
            output_path=OUTPUT_DIR / "metajudge_v0551_qa_audit_suspect.docx",
            show_suspect_reasons=True,
        )

    # Doc 4: Merged audit with justifications
    print("\nGenerating Merged Audit (with justifications)...")
    attach_justifications(all_items, justifications)
    generate_document(
        items=all_items,
        title="MetaJudge v0.5.5.1",
        subtitle="Complete QA Audit with Gold Answer Justifications",
        run_summary=run_summary,
        success_criteria=success_criteria,
        output_path=OUTPUT_DIR / "metajudge_v0551_qa_audit_with_justifications.docx",
    )

    print(f"\nDone! Documents saved to {OUTPUT_DIR}/")


if __name__ == "__main__":
    main()

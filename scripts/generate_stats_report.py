#!/usr/bin/env python3
"""
MetaJudge v0.5.5.1 — Statistical Analysis Report Generator

Produces three deliverables:
1. Stats theoretical backgrounder (markdown)
2. Polished results report (.docx with embedded figures)
3. Reproducibility log (markdown)

Usage:
    python scripts/generate_stats_report.py
"""

from __future__ import annotations

import csv
import hashlib
import json
import os
import platform
import sys
from collections import defaultdict
from itertools import combinations
from pathlib import Path

import numpy as np

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
PROJECT_ROOT = Path(__file__).resolve().parent.parent
DATA_DIR = Path("/tmp/v0551/v0.5.5.1 - results")
OUTPUT_DIR = PROJECT_ROOT / "outputs"
FIGURES_DIR = OUTPUT_DIR / "figures"

CAL_CSV = DATA_DIR / "calibration_item_audit (8).csv"
FB_CSV = DATA_DIR / "family_b_item_audit (8).csv"
RUN_SUMMARY = DATA_DIR / "run_summary (7).json"
BRIDGE_JSON = DATA_DIR / "bridge_report_all_models (2).json"
VERDICT_JSON = DATA_DIR / "success_criteria_verdict.json"
AUDIT_QUEUE_CSV = DATA_DIR / "audit_review_queue (2).csv"

# Short display names for models
MODEL_SHORT = {
    "google/gemini-2.5-flash": "Gemini Flash",
    "google/gemini-2.5-pro": "Gemini Pro",
    "anthropic/claude-sonnet-4@20250514": "Sonnet 4",
    "anthropic/claude-haiku-4-5@20251001": "Haiku 4.5",
    "deepseek-ai/deepseek-v3.1": "DeepSeek V3.1",
}

SEED = 42

# ---------------------------------------------------------------------------
# Data loading
# ---------------------------------------------------------------------------

def load_csv(path: Path) -> list[dict]:
    with open(path, newline="") as f:
        return list(csv.DictReader(f))


def load_json(path: Path) -> dict:
    with open(path) as f:
        return json.load(f)


def short_name(model: str) -> str:
    return MODEL_SHORT.get(model, model.split("/")[-1])


def load_all_data():
    """Load and structure all v0.5.5.1 data."""
    cal_rows = load_csv(CAL_CSV)
    fb_rows = load_csv(FB_CSV)
    run_summary = load_json(RUN_SUMMARY)
    audit_queue = load_csv(AUDIT_QUEUE_CSV)

    # Flagged item IDs for sensitivity analysis
    flagged_items = set(row["item_id"] for row in audit_queue)

    # Structure calibration data: {model: {item_id: {...}}}
    cal_by_model = defaultdict(dict)
    for row in cal_rows:
        m = row["model_name"]
        iid = row["item_id"]
        cal_by_model[m][iid] = {
            "is_correct": row["is_correct"] == "True",
            "brier_score": float(row["brier_score"]),
            "confidence": float(row["confidence"]),
            "mechanism": row["mechanism"],
            "gold_answer": row["gold_answer"],
            "model_answer": row["model_answer"],
        }

    # Structure Family B data: {model: {item_id: {...}}}
    fb_by_model = defaultdict(dict)
    for row in fb_rows:
        m = row["model_name"]
        iid = row["item_id"]
        fb_by_model[m][iid] = {
            "model_decision": row["model_decision"],
            "gold_action": row["gold_action"],
            "utility": float(row["utility"]),
            "confidence": float(row["confidence"]),
            "is_correct": row["is_correct"] == "True",
        }

    return {
        "cal_rows": cal_rows,
        "fb_rows": fb_rows,
        "cal_by_model": dict(cal_by_model),
        "fb_by_model": dict(fb_by_model),
        "run_summary": run_summary,
        "flagged_items": flagged_items,
    }


# ---------------------------------------------------------------------------
# Statistical computations (wrapping existing metajudge functions)
# ---------------------------------------------------------------------------

def compute_effect_sizes(scores_a, scores_b):
    """Cohen's d for paired differences."""
    a = np.asarray(scores_a, dtype=float)
    b = np.asarray(scores_b, dtype=float)
    diffs = a - b
    if np.std(diffs) == 0:
        return {"cohens_d": 0.0}
    d = float(np.mean(diffs) / np.std(diffs, ddof=1))
    return {"cohens_d": d}


def compute_odds_ratio(a_only: int, b_only: int) -> float:
    """Odds ratio from McNemar discordant cells."""
    if b_only == 0:
        return float("inf") if a_only > 0 else 1.0
    return a_only / b_only


def run_calibration_stats(cal_by_model, item_filter=None):
    """Run full pairwise calibration statistics."""
    from metajudge.scoring.statistics import (
        mcnemar_test,
        paired_permutation_test,
        paired_bootstrap_ci,
        holm_correction,
    )

    models = sorted(cal_by_model.keys())

    # Find common items
    common = set.intersection(*(set(cal_by_model[m].keys()) for m in models))
    if item_filter:
        common = common - item_filter
    item_ids = sorted(common)

    results = {}
    all_p = {}

    for m_a, m_b in combinations(models, 2):
        pair = f"{short_name(m_a)} vs {short_name(m_b)}"
        correct_a = [cal_by_model[m_a][i]["is_correct"] for i in item_ids]
        correct_b = [cal_by_model[m_b][i]["is_correct"] for i in item_ids]
        brier_a = [cal_by_model[m_a][i]["brier_score"] for i in item_ids]
        brier_b = [cal_by_model[m_b][i]["brier_score"] for i in item_ids]

        mcn = mcnemar_test(correct_a, correct_b)
        perm = paired_permutation_test(brier_a, brier_b)
        boot = paired_bootstrap_ci(brier_a, brier_b)
        eff = compute_effect_sizes(brier_a, brier_b)
        odds = compute_odds_ratio(mcn["a_only"], mcn["b_only"])

        results[pair] = {
            "n_items": len(item_ids),
            "acc_a": float(np.mean(correct_a)),
            "acc_b": float(np.mean(correct_b)),
            "mcnemar": mcn,
            "brier_mean_a": float(np.mean(brier_a)),
            "brier_mean_b": float(np.mean(brier_b)),
            "permutation": perm,
            "bootstrap_ci": boot,
            "cohens_d": eff["cohens_d"],
            "odds_ratio": odds,
        }
        all_p[f"{pair}_acc"] = mcn["p_value"]
        all_p[f"{pair}_brier"] = perm["p_value"]

    corrected = holm_correction(all_p)
    return {
        "models": models,
        "n_items": len(item_ids),
        "pairwise": results,
        "holm": corrected,
    }


def run_family_b_stats(fb_by_model, item_filter=None, min_items=10):
    """Run full pairwise Family B statistics."""
    from metajudge.scoring.statistics import (
        stuart_maxwell_test,
        paired_permutation_test,
        paired_bootstrap_ci,
        holm_correction,
    )

    # Exclude models with fewer than min_items
    models = sorted(m for m in fb_by_model if len(fb_by_model[m]) >= min_items)

    common = set.intersection(*(set(fb_by_model[m].keys()) for m in models))
    if item_filter:
        common = common - item_filter
    item_ids = sorted(common)

    results = {}
    all_p = {}

    for m_a, m_b in combinations(models, 2):
        pair = f"{short_name(m_a)} vs {short_name(m_b)}"
        actions_a = [fb_by_model[m_a][i]["model_decision"] for i in item_ids]
        actions_b = [fb_by_model[m_b][i]["model_decision"] for i in item_ids]
        util_a = [fb_by_model[m_a][i]["utility"] for i in item_ids]
        util_b = [fb_by_model[m_b][i]["utility"] for i in item_ids]

        sm = stuart_maxwell_test(actions_a, actions_b,
                                 labels=["answer", "clarify", "verify", "abstain"])
        perm = paired_permutation_test(util_a, util_b)
        boot = paired_bootstrap_ci(util_a, util_b)
        eff = compute_effect_sizes(util_a, util_b)

        results[pair] = {
            "n_items": len(item_ids),
            "util_mean_a": float(np.mean(util_a)),
            "util_mean_b": float(np.mean(util_b)),
            "stuart_maxwell": sm,
            "permutation": perm,
            "bootstrap_ci": boot,
            "cohens_d": eff["cohens_d"],
        }
        all_p[f"{pair}_util"] = perm["p_value"]
        if not np.isnan(sm["p_value"]):
            all_p[f"{pair}_actions"] = sm["p_value"]

    corrected = holm_correction(all_p) if all_p else {}
    return {
        "models": models,
        "n_items": len(item_ids),
        "pairwise": results,
        "holm": corrected,
    }


def compute_per_model_calibration_metrics(cal_by_model):
    """Compute ECE, overconfidence rate, accuracy per model."""
    from metajudge.scoring.calibration_metrics import (
        expected_calibration_error,
        overconfidence_rate,
        accuracy_by_confidence_bucket,
    )

    metrics = {}
    for model, items in cal_by_model.items():
        confs = [v["confidence"] for v in items.values()]
        corrects = [v["is_correct"] for v in items.values()]
        briers = [v["brier_score"] for v in items.values()]

        metrics[model] = {
            "accuracy": float(np.mean(corrects)),
            "mean_brier": float(np.mean(briers)),
            "ece": expected_calibration_error(confs, corrects),
            "overconfidence_rate": overconfidence_rate(confs, corrects),
            "n_items": len(items),
            "reliability_bins": accuracy_by_confidence_bucket(confs, corrects),
        }
    return metrics


def compute_per_model_fb_metrics(fb_by_model):
    """Compute UWAA, action accuracy, per-action F1 per model."""
    from metajudge.scoring.abstention_metrics import compute_uwaa, compute_action_f1

    metrics = {}
    for model, items in fb_by_model.items():
        utils = [v["utility"] for v in items.values()]
        decisions = [v["model_decision"] for v in items.values()]
        golds = [v["gold_action"] for v in items.values()]

        action_acc = float(np.mean([d == g for d, g in zip(decisions, golds)]))
        f1 = compute_action_f1(decisions, golds)

        metrics[model] = {
            "uwaa": compute_uwaa(utils),
            "mean_utility": float(np.mean(utils)),
            "action_accuracy": action_acc,
            "action_f1": f1,
            "n_items": len(items),
            "action_dist": {
                a: sum(1 for d in decisions if d == a) / len(decisions)
                for a in ["answer", "clarify", "verify", "abstain"]
            },
        }
    return metrics


def compute_per_mechanism_stats(cal_by_model):
    """Per-mechanism accuracy and Brier by model."""
    # Collect all mechanisms
    mechanisms = set()
    for items in cal_by_model.values():
        for v in items.values():
            mechanisms.add(v["mechanism"])
    mechanisms = sorted(mechanisms)

    models = sorted(cal_by_model.keys())
    result = {}
    for mech in mechanisms:
        result[mech] = {}
        for model in models:
            items = {k: v for k, v in cal_by_model[model].items()
                     if v["mechanism"] == mech}
            if not items:
                result[mech][model] = {"accuracy": float("nan"), "mean_brier": float("nan"), "n": 0}
                continue
            corrects = [v["is_correct"] for v in items.values()]
            briers = [v["brier_score"] for v in items.values()]
            result[mech][model] = {
                "accuracy": float(np.mean(corrects)),
                "mean_brier": float(np.mean(briers)),
                "n": len(items),
            }
    return result


def compute_bridge_correlations(cal_by_model):
    """Spearman correlation between confidence and correctness per model."""
    from metajudge.scoring.statistics import spearman_with_ci

    results = {}
    for model, items in cal_by_model.items():
        confs = [v["confidence"] for v in items.values()]
        corrects = [float(v["is_correct"]) for v in items.values()]
        results[model] = spearman_with_ci(confs, corrects)
    return results


# ---------------------------------------------------------------------------
# Placeholder for parts 2-5 (will be added incrementally)
# ---------------------------------------------------------------------------

def generate_visualizations(data, cal_stats, fb_stats, cal_metrics,
                            fb_metrics, mech_stats, bridge_corr,
                            cal_stats_clean, fb_stats_clean):
    """Generate all PNG visualizations."""
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import seaborn as sns

    sns.set_theme(style="whitegrid", font_scale=1.1)
    COLORS = sns.color_palette("tab10", n_colors=len(cal_metrics))
    model_order = sorted(cal_metrics.keys(), key=lambda m: short_name(m))
    model_colors = {m: COLORS[i] for i, m in enumerate(model_order)}

    # --- 1. Calibration reliability diagram ---
    fig, ax = plt.subplots(figsize=(8, 6))
    ax.plot([0, 1], [0, 1], "k--", alpha=0.5, label="Perfect calibration")
    for model in model_order:
        bins = cal_metrics[model]["reliability_bins"]
        confs = [b[0] for b in bins if b[2] > 0]
        accs = [b[1] for b in bins if b[2] > 0]
        ax.plot(confs, accs, "o-", color=model_colors[model],
                label=short_name(model), markersize=6)
    ax.set_xlabel("Mean Predicted Confidence")
    ax.set_ylabel("Observed Accuracy")
    ax.set_title("Calibration Reliability Diagram")
    ax.legend(loc="lower right", fontsize=9)
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    fig.tight_layout()
    fig.savefig(FIGURES_DIR / "cal_reliability_diagram.png", dpi=150)
    plt.close(fig)
    print("  [1/11] cal_reliability_diagram.png")

    # --- 2. Brier score forest plot ---
    pairs = list(cal_stats["pairwise"].keys())
    fig, ax = plt.subplots(figsize=(10, max(4, len(pairs) * 0.5)))
    y_pos = list(range(len(pairs)))
    for i, pair in enumerate(pairs):
        d = cal_stats["pairwise"][pair]
        ci = d["bootstrap_ci"]
        ax.errorbar(ci["observed_diff"], i,
                    xerr=[[ci["observed_diff"] - ci["ci_lower"]],
                          [ci["ci_upper"] - ci["observed_diff"]]],
                    fmt="o", color="steelblue", capsize=4, markersize=6)
    ax.axvline(0, color="red", linestyle="--", alpha=0.5)
    ax.set_yticks(y_pos)
    ax.set_yticklabels(pairs, fontsize=9)
    ax.set_xlabel("Brier Score Difference (A − B)")
    ax.set_title("Pairwise Brier Score Differences with 95% Bootstrap CI")
    fig.tight_layout()
    fig.savefig(FIGURES_DIR / "cal_brier_forest_plot.png", dpi=150)
    plt.close(fig)
    print("  [2/11] cal_brier_forest_plot.png")

    # --- 3. Per-mechanism accuracy heatmap ---
    mechanisms = sorted(mech_stats.keys())
    models_sorted = model_order
    acc_matrix = []
    for mech in mechanisms:
        row = [mech_stats[mech].get(m, {}).get("accuracy", float("nan"))
               for m in models_sorted]
        acc_matrix.append(row)
    acc_arr = np.array(acc_matrix)

    fig, ax = plt.subplots(figsize=(10, max(5, len(mechanisms) * 0.6)))
    sns.heatmap(acc_arr, annot=True, fmt=".2f", cmap="RdYlGn",
                xticklabels=[short_name(m) for m in models_sorted],
                yticklabels=mechanisms, ax=ax, vmin=0, vmax=1,
                linewidths=0.5)
    ax.set_title("Accuracy by Mechanism × Model")
    fig.tight_layout()
    fig.savefig(FIGURES_DIR / "cal_mechanism_heatmap.png", dpi=150)
    plt.close(fig)
    print("  [3/11] cal_mechanism_heatmap.png")

    # --- 4. Per-mechanism Brier bars ---
    fig, ax = plt.subplots(figsize=(12, 6))
    x = np.arange(len(mechanisms))
    width = 0.15
    for i, model in enumerate(models_sorted):
        vals = [mech_stats[mech].get(model, {}).get("mean_brier", 0)
                for mech in mechanisms]
        ax.bar(x + i * width, vals, width, label=short_name(model),
               color=model_colors[model])
    ax.set_xticks(x + width * (len(models_sorted) - 1) / 2)
    ax.set_xticklabels(mechanisms, rotation=45, ha="right", fontsize=9)
    ax.set_ylabel("Mean Brier Score")
    ax.set_title("Mean Brier Score by Mechanism × Model")
    ax.legend(fontsize=8)
    fig.tight_layout()
    fig.savefig(FIGURES_DIR / "cal_mechanism_brier_bars.png", dpi=150)
    plt.close(fig)
    print("  [4/11] cal_mechanism_brier_bars.png")

    # --- 5. Calibration p-value heatmap ---
    pair_names = list(cal_stats["pairwise"].keys())
    test_types = ["McNemar (acc)", "Permutation (Brier)"]
    pval_matrix = []
    for pair in pair_names:
        d = cal_stats["pairwise"][pair]
        pval_matrix.append([d["mcnemar"]["p_value"], d["permutation"]["p_value"]])
    pval_arr = np.array(pval_matrix)

    fig, ax = plt.subplots(figsize=(6, max(4, len(pair_names) * 0.5)))
    sns.heatmap(pval_arr, annot=True, fmt=".3f", cmap="RdYlGn_r",
                xticklabels=test_types, yticklabels=pair_names,
                ax=ax, vmin=0, vmax=0.2, linewidths=0.5)
    ax.set_title("Calibration P-values (darker = more significant)")
    fig.tight_layout()
    fig.savefig(FIGURES_DIR / "cal_pvalue_heatmap.png", dpi=150)
    plt.close(fig)
    print("  [5/11] cal_pvalue_heatmap.png")

    # --- 6. Family B action distribution ---
    fb_model_order = sorted(fb_metrics.keys(), key=lambda m: short_name(m))
    actions = ["answer", "clarify", "verify", "abstain"]
    action_colors = {"answer": "#4CAF50", "clarify": "#2196F3",
                     "verify": "#FF9800", "abstain": "#F44336"}

    fig, ax = plt.subplots(figsize=(10, 6))
    bottom = np.zeros(len(fb_model_order))
    for action in actions:
        vals = [fb_metrics[m]["action_dist"].get(action, 0) for m in fb_model_order]
        ax.bar([short_name(m) for m in fb_model_order], vals, bottom=bottom,
               label=action, color=action_colors[action])
        bottom += np.array(vals)
    ax.set_ylabel("Proportion")
    ax.set_title("Action Distribution by Model (Family B)")
    ax.legend()
    fig.tight_layout()
    fig.savefig(FIGURES_DIR / "fb_action_distribution.png", dpi=150)
    plt.close(fig)
    print("  [6/11] fb_action_distribution.png")

    # --- 7. Family B action confusion heatmaps ---
    # One subplot per model (exclude models with <10 items)
    fb_models_full = [m for m in fb_model_order if fb_metrics[m]["n_items"] >= 10]
    n_models = len(fb_models_full)
    fig, axes = plt.subplots(1, n_models, figsize=(4 * n_models, 4))
    if n_models == 1:
        axes = [axes]
    for idx, model in enumerate(fb_models_full):
        items = data["fb_by_model"][model]
        conf_mat = np.zeros((4, 4), dtype=int)
        for v in items.values():
            gi = actions.index(v["gold_action"]) if v["gold_action"] in actions else -1
            di = actions.index(v["model_decision"]) if v["model_decision"] in actions else -1
            if gi >= 0 and di >= 0:
                conf_mat[gi][di] += 1
        sns.heatmap(conf_mat, annot=True, fmt="d", cmap="Blues",
                    xticklabels=actions, yticklabels=actions,
                    ax=axes[idx], linewidths=0.5)
        axes[idx].set_title(short_name(model), fontsize=10)
        axes[idx].set_xlabel("Model Decision")
        axes[idx].set_ylabel("Gold Action")
    fig.suptitle("Action Confusion Matrices (Family B)", fontsize=13)
    fig.tight_layout()
    fig.savefig(FIGURES_DIR / "fb_action_confusion.png", dpi=150)
    plt.close(fig)
    print("  [7/11] fb_action_confusion.png")

    # --- 8. Family B utility forest plot ---
    if fb_stats["pairwise"]:
        fb_pairs = list(fb_stats["pairwise"].keys())
        fig, ax = plt.subplots(figsize=(10, max(4, len(fb_pairs) * 0.5)))
        for i, pair in enumerate(fb_pairs):
            d = fb_stats["pairwise"][pair]
            ci = d["bootstrap_ci"]
            ax.errorbar(ci["observed_diff"], i,
                        xerr=[[ci["observed_diff"] - ci["ci_lower"]],
                              [ci["ci_upper"] - ci["observed_diff"]]],
                        fmt="o", color="darkorange", capsize=4, markersize=6)
        ax.axvline(0, color="red", linestyle="--", alpha=0.5)
        ax.set_yticks(range(len(fb_pairs)))
        ax.set_yticklabels(fb_pairs, fontsize=9)
        ax.set_xlabel("Utility Difference (A − B)")
        ax.set_title("Pairwise Utility Differences with 95% Bootstrap CI (Family B)")
        fig.tight_layout()
        fig.savefig(FIGURES_DIR / "fb_utility_forest_plot.png", dpi=150)
        plt.close(fig)
    print("  [8/11] fb_utility_forest_plot.png")

    # --- 9. Confidence distribution violin plots ---
    conf_data = []
    conf_labels = []
    for model in model_order:
        confs = [v["confidence"] for v in data["cal_by_model"][model].values()]
        conf_data.extend(confs)
        conf_labels.extend([short_name(model)] * len(confs))

    fig, ax = plt.subplots(figsize=(10, 6))
    import pandas as pd
    df_conf = pd.DataFrame({"Confidence": conf_data, "Model": conf_labels})
    sns.violinplot(data=df_conf, x="Model", y="Confidence", ax=ax,
                   palette="tab10", inner="quartile")
    ax.set_title("Confidence Distribution by Model (Calibration)")
    ax.set_ylim(0, 1.05)
    fig.tight_layout()
    fig.savefig(FIGURES_DIR / "bridge_confidence_violins.png", dpi=150)
    plt.close(fig)
    print("  [9/11] bridge_confidence_violins.png")

    # --- 10. Confidence-accuracy scatter with Spearman ---
    fig, ax = plt.subplots(figsize=(8, 6))
    for model in model_order:
        bins = cal_metrics[model]["reliability_bins"]
        confs = [b[0] for b in bins if b[2] > 0]
        accs = [b[1] for b in bins if b[2] > 0]
        rho = bridge_corr[model]["rho"]
        ax.plot(confs, accs, "o-", color=model_colors[model],
                label=f"{short_name(model)} (ρ={rho:.2f})", markersize=5)
    ax.plot([0, 1], [0, 1], "k--", alpha=0.3)
    ax.set_xlabel("Confidence")
    ax.set_ylabel("Accuracy")
    ax.set_title("Confidence vs Accuracy with Spearman ρ")
    ax.legend(fontsize=8, loc="lower right")
    fig.tight_layout()
    fig.savefig(FIGURES_DIR / "bridge_confidence_accuracy.png", dpi=150)
    plt.close(fig)
    print("  [10/11] bridge_confidence_accuracy.png")

    # --- 11. Master significance heatmap ---
    all_pairs = list(cal_stats["pairwise"].keys())
    # Combine cal and FB test p-values
    test_cols = ["McNemar", "Brier Perm", "Utility Perm", "Stuart-Maxwell"]
    sig_matrix = []
    for pair in all_pairs:
        row = []
        cd = cal_stats["pairwise"].get(pair, {})
        row.append(cd.get("mcnemar", {}).get("p_value", float("nan")))
        row.append(cd.get("permutation", {}).get("p_value", float("nan")))
        fd = fb_stats["pairwise"].get(pair, {})
        row.append(fd.get("permutation", {}).get("p_value", float("nan")))
        row.append(fd.get("stuart_maxwell", {}).get("p_value", float("nan")))
        sig_matrix.append(row)
    sig_arr = np.array(sig_matrix)

    fig, ax = plt.subplots(figsize=(8, max(4, len(all_pairs) * 0.5)))
    sns.heatmap(sig_arr, annot=True, fmt=".3f", cmap="RdYlGn_r",
                xticklabels=test_cols, yticklabels=all_pairs,
                ax=ax, vmin=0, vmax=0.2, linewidths=0.5,
                mask=np.isnan(sig_arr))
    ax.set_title("Master Significance Heatmap (p-values)")
    fig.tight_layout()
    fig.savefig(FIGURES_DIR / "significance_master_heatmap.png", dpi=150)
    plt.close(fig)
    print("  [11/11] significance_master_heatmap.png")


def generate_backgrounder(output_path):
    """Generate stats theoretical backgrounder markdown."""
    text = """# MetaJudge v0.5.5.1 — Statistical Testing Backgrounder

## 1. Introduction

MetaJudge is a benchmark measuring metacognitive behavior in large language models.
Statistical comparisons between models must account for three structural features
of the data: (a) observations are **paired** — the same items are administered to
every model, creating natural within-item dependencies; (b) sample sizes are
**moderate** (n = 117 calibration items, n = 84 selective-abstention items),
placing the analysis in a regime where large-sample approximations are not always
reliable; and (c) the primary scoring metrics (Brier scores, utility payoffs) have
**non-normal, bounded distributions** that discourage parametric assumptions.

For these reasons, MetaJudge's statistical framework is **non-parametric by
default**, relying on exact tests, permutation procedures, and bootstrap
resampling rather than t-tests or ANOVA. All pairwise comparisons are corrected
for multiple testing using familywise error-rate control.

---

## 2. Statistical Tests

### 2.1 McNemar's Test

McNemar's test (McNemar, 1947) is the standard method for comparing two
classifiers evaluated on the same test set when the outcome is binary
(correct / incorrect). It considers only the **discordant pairs** — items where
one model is correct and the other is wrong — and tests whether the two
directions of disagreement are equally likely under the null hypothesis.

**Why chosen for MetaJudge:** Each calibration item yields a binary correctness
judgment per model, and all models see the same items. McNemar's test is the
natural paired test for this structure, recommended by Dietterich (1998) as the
preferred statistical test for comparing supervised classifiers.

**Implementation detail:** For small discordant counts (n < 25), the exact
binomial test is used; for larger counts, the chi-squared approximation with
continuity correction is applied.

**Leading alternatives considered:**
- *Cochran's Q test* — extends McNemar's to simultaneous comparison of k > 2
  classifiers (Cochran, 1950). Rejected because it provides only an omnibus
  "any difference" verdict, losing pairwise granularity. It may be added as a
  gatekeeper test when the model count exceeds 8.
- *Exact conditional test (Edwards, 1948)* — provides exact p-values without
  continuity correction but is computationally heavier and unnecessary given
  the binomial fallback already implemented.

### 2.2 Paired Permutation Test

The paired permutation test (Good, 2005; Pesarin & Salmaso, 2010) is a
distribution-free hypothesis test for the difference in means of two paired
samples. Under the null hypothesis that the two score distributions are
exchangeable, the test randomly flips the sign of each paired difference and
recomputes the test statistic across many iterations, constructing an empirical
null distribution.

**Why chosen for MetaJudge:** Brier scores and utility values are continuous but
bounded on [0, 1] and [-1, 1] respectively, with distributions that are
typically skewed and zero-inflated. The permutation test makes no distributional
assumptions beyond exchangeability under the null, making it strictly more
general than the Wilcoxon signed-rank test.

**Implementation detail:** 10,000 permutations are used with a fixed random
seed (42) for reproducibility. The two-sided p-value is computed as
(count + 1) / (n_perm + 1), following the recommendation of Phipson and Smyth
(2010) to avoid zero p-values.

**Leading alternatives considered:**
- *Wilcoxon signed-rank test* — a classical non-parametric paired test
  (Wilcoxon, 1945). Rejected because it assumes a symmetric distribution of
  differences, which Brier score differences often violate.
- *Paired t-test* — parametric, assumes normal differences. Rejected due to
  bounded, skewed score distributions.

### 2.3 Bootstrap Confidence Intervals

The percentile bootstrap (Efron, 1979; Efron & Tibshirani, 1993) constructs
confidence intervals by resampling with replacement from the observed data.
For paired comparisons, the item-level differences are resampled; for
single-sample estimates, the raw scores are resampled. The α/2 and
1 − α/2 percentiles of the bootstrap distribution define the interval.

**Why chosen for MetaJudge:** Bootstrap CIs complement hypothesis tests by
quantifying the **magnitude and uncertainty** of effects, not just statistical
significance. They require no distributional assumptions and are straightforward
to implement for any summary statistic.

**Implementation detail:** 10,000 bootstrap resamples, seed = 42, percentile
method. Both paired-difference CIs and single-sample CIs are provided.

**Leading alternatives considered:**
- *BCa (bias-corrected and accelerated) bootstrap* — provides higher-order
  accuracy for skewed distributions (Efron, 1987). Rejected because the added
  complexity yields marginal benefit at n = 117, and the percentile method is
  more transparent.
- *Bayesian credible intervals* — would require specifying priors for Brier
  scores. Rejected to maintain a frequentist framework consistent with the
  other tests.

### 2.4 Stuart-Maxwell Test

The Stuart-Maxwell test (Stuart, 1955; Maxwell, 1970) is an extension of
McNemar's test to square contingency tables with k > 2 categories. It tests
marginal homogeneity — whether the row and column marginals of a paired
k × k table are equal — which corresponds to asking whether two models have
the same distribution of predicted actions.

**Why chosen for MetaJudge:** Family B items require models to choose among
four actions (answer, clarify, verify, abstain). Comparing the action
distributions of two models on the same items requires a paired multi-category
test. The Stuart-Maxwell test is the standard choice.

**Implementation detail:** The test statistic follows a chi-squared distribution
with k − 1 degrees of freedom under the null. For singular or near-singular
variance matrices, the Moore-Penrose pseudoinverse is used.

**Leading alternatives considered:**
- *Bhapkar test* — asymptotically equivalent to Stuart-Maxwell but uses a
  different variance estimator (Bhapkar, 1966). Less widely known; rejected
  for familiarity and interpretability.
- *McNemar-Bowker test of symmetry* — tests whether the off-diagonal cells
  are symmetric, which is a stronger hypothesis than marginal homogeneity.
  Rejected because we specifically want to compare marginals, not symmetry.

### 2.5 Spearman Rank Correlation

Spearman's rank correlation coefficient (Spearman, 1904) measures the monotonic
association between two variables by computing the Pearson correlation on their
ranks. It is non-parametric and robust to outliers and non-linear relationships.

**Why chosen for MetaJudge:** The bridge analysis examines whether higher model
confidence is associated with higher accuracy. This is a monotonic relationship
question (not necessarily linear), making Spearman's ρ the natural choice.
Bootstrap CIs are provided for inferential precision.

**Leading alternatives considered:**
- *Kendall's τ* — another rank correlation that is more robust to ties and has
  a simpler variance formula (Kendall, 1938). Rejected because Spearman's ρ
  has greater statistical power for continuous data without excessive ties.
- *Pearson's r* — assumes linear relationship and bivariate normality. Rejected
  because confidence-accuracy relationships are typically non-linear.

### 2.6 Holm-Bonferroni Correction

The Holm-Bonferroni procedure (Holm, 1979) is a sequentially rejective method
for controlling the familywise error rate (FWER) when conducting multiple
hypothesis tests. It orders p-values from smallest to largest and compares each
to α / (m − rank + 1), stopping at the first non-rejection.

**Why chosen for MetaJudge:** With 5 models producing 10 pairwise comparisons
and multiple test types per pair, the raw number of hypothesis tests ranges from
20 to 40+. Without correction, the probability of at least one false positive
becomes unacceptably high. Holm-Bonferroni is uniformly more powerful than the
classical Bonferroni correction while maintaining strong FWER control.

**Leading alternatives considered:**
- *Benjamini-Hochberg (BH) procedure* — controls the false discovery rate
  (FDR), which is less conservative than FWER (Benjamini & Hochberg, 1995).
  Rejected for the current 5-model analysis because FWER control is preferred
  for benchmark validation claims. BH may be adopted when model count exceeds
  8, as Holm becomes overly conservative at 100+ comparisons.
- *Bonferroni correction* — the simplest FWER method. Rejected because Holm
  is uniformly more powerful with no added complexity.

---

## 3. Scoring Metrics

### 3.1 Brier Score

The Brier score (Brier, 1950) is a strictly proper scoring rule that measures
the accuracy of probabilistic predictions. For a single item with stated
confidence c and binary outcome y ∈ {0, 1}:

    Brier = (c − y)²

A strictly proper scoring rule has the unique property that the expected score
is minimized (for lower-is-better formulations) when the stated confidence
equals the true probability of being correct (Gneiting & Raftery, 2007). This
makes it impossible to systematically improve one's Brier score by misreporting
confidence — the optimal strategy is honest calibration.

MetaJudge reports the affine transform 1 − Brier (higher-is-better, range
[0, 1]) in leaderboards, but statistical tests operate on the raw Brier loss.

### 3.2 Expected Calibration Error (ECE)

Expected Calibration Error (Naeini, Cooper, & Hauskrecht, 2015) partitions
predictions into bins by confidence level and measures the weighted average
gap between confidence and accuracy within each bin:

    ECE = Σ (|B_k| / N) · |acc(B_k) − conf(B_k)|

ECE is used as a **diagnostic metric** in MetaJudge (not in scoring) because
it captures a complementary aspect of calibration quality that Brier scores
can obscure — specifically, the systematic direction and pattern of
miscalibration across the confidence spectrum.

### 3.3 Utility-Weighted Action Accuracy (UWAA)

UWAA is a project-specific metric for Family B (selective abstention) that
normalizes the mean per-item utility to [0, 1]:

    UWAA = (mean_utility + 1.0) / 2.0

The utility for each item is determined by a 5 × 4 payoff matrix that maps
(model_decision × gold_action) pairs to scalar rewards in [-1, +1]. The
matrix is designed so that correct answers to answerable questions yield +1.0,
incorrect answers to answerable questions yield -1.0, and appropriate
deferral actions (clarify, verify, abstain) yield partial credit.

---

## 4. Effect Size Measures

### 4.1 Cohen's d (Paired)

For continuous paired differences (Brier scores, utility):

    d = mean(A − B) / sd(A − B)

Conventional thresholds: |d| < 0.2 negligible, 0.2–0.5 small, 0.5–0.8
medium, > 0.8 large (Cohen, 1988). These thresholds are guidelines, not
absolute cutoffs.

### 4.2 McNemar Odds Ratio

For binary accuracy comparisons, the odds ratio of discordant cells:

    OR = (A correct, B wrong) / (B correct, A wrong)

An OR > 1 indicates model A is more accurate; OR < 1 favors model B.

---

## 5. Power and Limitations

**Sample sizes:** n = 117 calibration items, n = 84 Family B items. With
5 models, there are 10 pairwise comparisons for calibration and 6 for
Family B (Gemini Flash excluded due to n = 1 in Family B).

**Power considerations:**
- McNemar's test: power depends on the number of discordant pairs, not total n.
  With 117 items, typical discordant counts range from 15-40 pairs, which
  provides adequate power for detecting accuracy differences of ~10 percentage
  points.
- Permutation tests at n = 117: sufficient to detect Brier score differences
  with Cohen's d ≥ 0.3 at 80% power.
- Per-mechanism subgroups (n = 2 to n = 20) are reported as **descriptive
  only** — formal hypothesis testing within mechanisms is underpowered.

**Sensitivity analysis:** All pairwise tests are run on both the full item set
and a clean subset excluding ~48 items flagged in the audit review queue. This
characterizes the robustness of findings to potential item quality concerns.

**Multiple comparisons:** With 10 pairwise comparisons × 2 tests per pair
(accuracy + Brier) = 20 tests, Holm-Bonferroni correction is applied. The
effective significance threshold for the most significant test is
α / 20 = 0.0025.

---

## 6. References

Benjamini, Y., & Hochberg, Y. (1995). Controlling the false discovery rate:
A practical and powerful approach to multiple testing. *Journal of the Royal
Statistical Society: Series B*, 57(1), 289–300.

Bhapkar, V. P. (1966). A note on the equivalence of two test criteria for
hypotheses in categorical data. *Journal of the American Statistical
Association*, 61(313), 228–235.

Brier, G. W. (1950). Verification of forecasts expressed in terms of
probability. *Monthly Weather Review*, 78(1), 1–3.

Cochran, W. G. (1950). The comparison of percentages in matched samples.
*Biometrika*, 37(3/4), 256–266.

Cohen, J. (1988). *Statistical Power Analysis for the Behavioral Sciences*
(2nd ed.). Lawrence Erlbaum Associates.

Dietterich, T. G. (1998). Approximate statistical tests for comparing
supervised classification learning algorithms. *Neural Computation*, 10(7),
1895–1923.

Efron, B. (1979). Bootstrap methods: Another look at the jackknife.
*The Annals of Statistics*, 7(1), 1–26.

Efron, B. (1987). Better bootstrap confidence intervals. *Journal of the
American Statistical Association*, 82(397), 171–185.

Efron, B., & Tibshirani, R. J. (1993). *An Introduction to the Bootstrap*.
Chapman & Hall/CRC.

Gneiting, T., & Raftery, A. E. (2007). Strictly proper scoring rules,
prediction, and estimation. *Journal of the American Statistical Association*,
102(477), 359–378.

Good, P. I. (2005). *Permutation, Parametric, and Bootstrap Tests of
Hypotheses* (3rd ed.). Springer.

Holm, S. (1979). A simple sequentially rejective multiple test procedure.
*Scandinavian Journal of Statistics*, 6(2), 65–70.

Kendall, M. G. (1938). A new measure of rank correlation. *Biometrika*,
30(1/2), 81–93.

Maxwell, A. E. (1970). Comparing the classification of subjects by two
independent judges. *British Journal of Psychiatry*, 116(535), 651–655.

McNemar, Q. (1947). Note on the sampling error of the difference between
correlated proportions or percentages. *Psychometrika*, 12(2), 153–157.

Naeini, M. P., Cooper, G., & Hauskrecht, M. (2015). Obtaining well-calibrated
probabilities using Bayesian binning. *Proceedings of the AAAI Conference on
Artificial Intelligence*, 29(1), 2901–2907.

Pesarin, F., & Salmaso, L. (2010). *Permutation Tests for Complex Data*.
Wiley.

Phipson, B., & Smyth, G. K. (2010). Permutation P-values should never be
zero. *Statistical Applications in Genetics and Molecular Biology*, 9(1).

Spearman, C. (1904). The proof and measurement of association between two
things. *American Journal of Psychology*, 15(1), 72–101.

Stuart, A. (1955). A test for homogeneity of the marginal distributions in a
two-way classification. *Biometrika*, 42(3/4), 412–416.

Wilcoxon, F. (1945). Individual comparisons by ranking methods. *Biometrics
Bulletin*, 1(6), 80–83.
"""
    with open(output_path, "w") as f:
        f.write(text.strip() + "\n")
    print(f"  Backgrounder: {output_path}")


def generate_docx_report(output_path, data, cal_stats, fb_stats,
                         cal_metrics, fb_metrics, mech_stats, bridge_corr,
                         cal_stats_clean, fb_stats_clean):
    """Generate polished .docx results report with embedded figures."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()

    # -- Styles --
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def add_heading(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        return h

    def add_table_from_data(headers, rows):
        """Add a formatted table to the document."""
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Light Grid Accent 1"
        # Header row
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(9)
        # Data rows
        for r_idx, row in enumerate(rows):
            for c_idx, val in enumerate(row):
                cell = table.rows[r_idx + 1].cells[c_idx]
                cell.text = str(val)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)
        return table

    def fmt_p(val):
        """Format p-value for display."""
        if val < 0.001:
            return "<0.001"
        return f"{val:.3f}"

    def fmt_f(val, decimals=3):
        """Format float."""
        if isinstance(val, float) and (np.isnan(val) or np.isinf(val)):
            return "N/A"
        return f"{val:.{decimals}f}"

    def sig_star(p):
        if p < 0.001:
            return "***"
        if p < 0.01:
            return "**"
        if p < 0.05:
            return "*"
        return ""

    # ===== TITLE PAGE =====
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("MetaJudge v0.5.5.1\nStatistical Analysis Report")
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Preliminary — Full Item Set + Sensitivity Analysis")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"Models: {len(cal_stats['models'])} | "
                 f"Calibration items: {cal_stats['n_items']} | "
                 f"Family B items: {fb_stats['n_items']}\n"
                 f"Date: 2026-03-28 | Benchmark version: v0.5.5.1").font.size = Pt(10)

    doc.add_page_break()

    # ===== 1. EXECUTIVE SUMMARY =====
    add_heading("1. Executive Summary", level=1)

    doc.add_paragraph(
        "This report presents pairwise statistical comparisons across five frontier "
        "language models on the MetaJudge v0.5.5.1 benchmark, covering both "
        "confidence calibration (Family A, 117 items) and selective abstention "
        "(Family B, 84 items). All tests are non-parametric with Holm-Bonferroni "
        "correction for multiple comparisons."
    )

    # Count significant results
    n_sig_cal = sum(1 for v in cal_stats["holm"].values()
                    if v.get("significant_0.05"))
    n_sig_fb = sum(1 for v in fb_stats["holm"].values()
                   if v.get("significant_0.05"))
    total_tests = len(cal_stats["holm"]) + len(fb_stats["holm"])

    doc.add_paragraph(
        f"Key findings: {n_sig_cal + n_sig_fb} of {total_tests} tests reached "
        f"statistical significance after Holm-Bonferroni correction "
        f"(α = 0.05). {n_sig_cal} in calibration, {n_sig_fb} in Family B."
    )

    doc.add_page_break()

    # ===== 2. FAMILY A: CALIBRATION RESULTS =====
    add_heading("2. Family A: Calibration Results", level=1)

    # 2.1 Model leaderboard
    add_heading("2.1 Model Leaderboard", level=2)
    model_order = sorted(cal_metrics.keys(), key=lambda m: short_name(m))
    headers = ["Model", "n", "Accuracy", "Mean Brier", "ECE", "Overconf. Rate"]
    rows = []
    for m in model_order:
        d = cal_metrics[m]
        rows.append([
            short_name(m), str(d["n_items"]),
            fmt_f(d["accuracy"]), fmt_f(d["mean_brier"]),
            fmt_f(d["ece"]), fmt_f(d["overconfidence_rate"]),
        ])
    add_table_from_data(headers, rows)

    # 2.2 Pairwise comparisons
    add_heading("2.2 Pairwise Statistical Comparisons", level=2)
    doc.add_paragraph(
        "McNemar's test compares binary accuracy; paired permutation test compares "
        "mean Brier scores. Bootstrap 95% CIs quantify effect magnitude. "
        "Cohen's d measures standardized effect size."
    )

    headers = ["Pair", "Acc A", "Acc B", "McNemar p", "Brier Δ", "Perm p",
               "95% CI", "Cohen's d", "Sig"]
    rows = []
    for pair, d in cal_stats["pairwise"].items():
        ci = d["bootstrap_ci"]
        # Check if any test for this pair is significant after correction
        holm_acc = cal_stats["holm"].get(f"{pair}_acc", {})
        holm_brier = cal_stats["holm"].get(f"{pair}_brier", {})
        is_sig = (holm_acc.get("significant_0.05", False) or
                  holm_brier.get("significant_0.05", False))
        rows.append([
            pair,
            fmt_f(d["acc_a"]), fmt_f(d["acc_b"]),
            fmt_p(d["mcnemar"]["p_value"]),
            fmt_f(ci["observed_diff"]),
            fmt_p(d["permutation"]["p_value"]),
            f"[{fmt_f(ci['ci_lower'])}, {fmt_f(ci['ci_upper'])}]",
            fmt_f(d["cohens_d"], 2),
            "Yes" + sig_star(min(d["mcnemar"]["p_value"],
                                 d["permutation"]["p_value"])) if is_sig else "No",
        ])
    add_table_from_data(headers, rows)

    # 2.3 Figures
    add_heading("2.3 Calibration Reliability Diagram", level=2)
    doc.add_picture(str(FIGURES_DIR / "cal_reliability_diagram.png"),
                    width=Inches(5.5))

    add_heading("2.4 Brier Score Forest Plot", level=2)
    doc.add_picture(str(FIGURES_DIR / "cal_brier_forest_plot.png"),
                    width=Inches(5.5))

    add_heading("2.5 Per-Mechanism Analysis", level=2)
    doc.add_picture(str(FIGURES_DIR / "cal_mechanism_heatmap.png"),
                    width=Inches(5.5))
    doc.add_paragraph()
    doc.add_picture(str(FIGURES_DIR / "cal_mechanism_brier_bars.png"),
                    width=Inches(5.5))

    # Mechanism summary table
    mechanisms = sorted(mech_stats.keys())
    headers = ["Mechanism"] + [short_name(m) for m in model_order] + ["n"]
    rows = []
    for mech in mechanisms:
        row = [mech]
        n_items = 0
        for m in model_order:
            md = mech_stats[mech].get(m, {})
            row.append(fmt_f(md.get("accuracy", float("nan"))))
            n_items = max(n_items, md.get("n", 0))
        underpowered = " ⚠" if n_items < 10 else ""
        row.append(f"{n_items}{underpowered}")
        rows.append(row)
    add_table_from_data(headers, rows)
    doc.add_paragraph("⚠ = fewer than 10 items (descriptive only)", style="Normal")

    add_heading("2.6 P-value Heatmap", level=2)
    doc.add_picture(str(FIGURES_DIR / "cal_pvalue_heatmap.png"),
                    width=Inches(4.5))

    doc.add_page_break()

    # ===== 3. FAMILY B: SELECTIVE ABSTENTION =====
    add_heading("3. Family B: Selective Abstention Results", level=1)

    # 3.1 Model leaderboard
    add_heading("3.1 Model Leaderboard", level=2)
    fb_model_order = sorted(fb_metrics.keys(), key=lambda m: short_name(m))
    headers = ["Model", "n", "UWAA", "Mean Utility", "Action Acc", "Macro F1"]
    rows = []
    for m in fb_model_order:
        d = fb_metrics[m]
        rows.append([
            short_name(m), str(d["n_items"]),
            fmt_f(d["uwaa"]), fmt_f(d["mean_utility"]),
            fmt_f(d["action_accuracy"]),
            fmt_f(d["action_f1"]["macro"]["f1"]),
        ])
    add_table_from_data(headers, rows)

    if fb_metrics.get(list(fb_metrics.keys())[0], {}).get("n_items", 0) == 1:
        doc.add_paragraph(
            "Note: Gemini Flash has only 1 Family B item and is excluded from "
            "pairwise statistical comparisons."
        )

    # 3.2 Pairwise comparisons
    if fb_stats["pairwise"]:
        add_heading("3.2 Pairwise Statistical Comparisons", level=2)
        headers = ["Pair", "Util A", "Util B", "Util Perm p", "S-M p",
                   "95% CI", "Cohen's d", "Sig"]
        rows = []
        for pair, d in fb_stats["pairwise"].items():
            ci = d["bootstrap_ci"]
            holm_util = fb_stats["holm"].get(f"{pair}_util", {})
            holm_act = fb_stats["holm"].get(f"{pair}_actions", {})
            is_sig = (holm_util.get("significant_0.05", False) or
                      holm_act.get("significant_0.05", False))
            rows.append([
                pair,
                fmt_f(d["util_mean_a"]), fmt_f(d["util_mean_b"]),
                fmt_p(d["permutation"]["p_value"]),
                fmt_p(d["stuart_maxwell"]["p_value"]),
                f"[{fmt_f(ci['ci_lower'])}, {fmt_f(ci['ci_upper'])}]",
                fmt_f(d["cohens_d"], 2),
                "Yes" + sig_star(min(d["permutation"]["p_value"],
                                     d["stuart_maxwell"]["p_value"])) if is_sig else "No",
            ])
        add_table_from_data(headers, rows)

    # 3.3 Figures
    add_heading("3.3 Action Distribution", level=2)
    doc.add_picture(str(FIGURES_DIR / "fb_action_distribution.png"),
                    width=Inches(5.5))

    add_heading("3.4 Action Confusion Matrices", level=2)
    doc.add_picture(str(FIGURES_DIR / "fb_action_confusion.png"),
                    width=Inches(6.0))

    if (FIGURES_DIR / "fb_utility_forest_plot.png").exists():
        add_heading("3.5 Utility Forest Plot", level=2)
        doc.add_picture(str(FIGURES_DIR / "fb_utility_forest_plot.png"),
                        width=Inches(5.5))

    doc.add_page_break()

    # ===== 4. BRIDGE ANALYSIS =====
    add_heading("4. Cross-Family Bridge Analysis", level=1)

    add_heading("4.1 Confidence-Accuracy Correlation", level=2)
    headers = ["Model", "Spearman ρ", "p-value", "95% CI"]
    rows = []
    for m in model_order:
        d = bridge_corr[m]
        rows.append([
            short_name(m), fmt_f(d["rho"], 3), fmt_p(d["p_value"]),
            f"[{fmt_f(d['ci_lower'], 3)}, {fmt_f(d['ci_upper'], 3)}]",
        ])
    add_table_from_data(headers, rows)

    doc.add_picture(str(FIGURES_DIR / "bridge_confidence_accuracy.png"),
                    width=Inches(5.5))

    add_heading("4.2 Confidence Distributions", level=2)
    doc.add_picture(str(FIGURES_DIR / "bridge_confidence_violins.png"),
                    width=Inches(5.5))

    doc.add_page_break()

    # ===== 5. SIGNIFICANCE SUMMARY =====
    add_heading("5. Statistical Significance Summary", level=1)

    add_heading("5.1 Master P-value Table (Holm-Corrected)", level=2)
    all_holm = {**cal_stats["holm"], **fb_stats["holm"]}
    headers = ["Test", "Raw p", "Adjusted p", "Significant"]
    rows = []
    for name, d in sorted(all_holm.items(), key=lambda x: x[1]["raw_p"]):
        rows.append([
            name, fmt_p(d["raw_p"]), fmt_p(d["adjusted_p"]),
            "Yes" + sig_star(d["adjusted_p"]) if d["significant_0.05"] else "No",
        ])
    add_table_from_data(headers, rows)

    add_heading("5.2 Master Significance Heatmap", level=2)
    doc.add_picture(str(FIGURES_DIR / "significance_master_heatmap.png"),
                    width=Inches(5.5))

    doc.add_page_break()

    # ===== 6. SENSITIVITY ANALYSIS =====
    add_heading("6. Robustness to Item Exclusion", level=1)

    doc.add_paragraph(
        f"Sensitivity analysis: {len(data['flagged_items'])} items flagged in the "
        f"audit review queue were excluded to assess robustness. This reduced "
        f"calibration from {cal_stats['n_items']} to {cal_stats_clean['n_items']} "
        f"items and Family B from {fb_stats['n_items']} to {fb_stats_clean['n_items']} items."
    )

    add_heading("6.1 Calibration — Full vs Clean", level=2)
    headers = ["Pair", "Full Brier p", "Clean Brier p", "Full Sig", "Clean Sig"]
    rows = []
    for pair in cal_stats["pairwise"]:
        full_p = cal_stats["pairwise"][pair]["permutation"]["p_value"]
        clean_d = cal_stats_clean["pairwise"].get(pair, {})
        clean_p = clean_d.get("permutation", {}).get("p_value", float("nan"))
        full_sig = cal_stats["holm"].get(f"{pair}_brier", {}).get("significant_0.05", False)
        clean_sig = cal_stats_clean["holm"].get(f"{pair}_brier", {}).get("significant_0.05", False)
        rows.append([
            pair, fmt_p(full_p), fmt_p(clean_p),
            "Yes" if full_sig else "No",
            "Yes" if clean_sig else "No",
        ])
    add_table_from_data(headers, rows)

    if fb_stats_clean["pairwise"]:
        add_heading("6.2 Family B — Full vs Clean", level=2)
        headers = ["Pair", "Full Util p", "Clean Util p", "Full Sig", "Clean Sig"]
        rows = []
        for pair in fb_stats["pairwise"]:
            full_p = fb_stats["pairwise"][pair]["permutation"]["p_value"]
            clean_d = fb_stats_clean["pairwise"].get(pair, {})
            clean_p = clean_d.get("permutation", {}).get("p_value", float("nan"))
            full_sig = fb_stats["holm"].get(f"{pair}_util", {}).get("significant_0.05", False)
            clean_sig = fb_stats_clean["holm"].get(f"{pair}_util", {}).get("significant_0.05", False)
            rows.append([
                pair, fmt_p(full_p), fmt_p(clean_p),
                "Yes" if full_sig else "No",
                "Yes" if clean_sig else "No",
            ])
        add_table_from_data(headers, rows)

    # Save
    doc.save(str(output_path))
    print(f"  Report: {output_path}")


def generate_reproducibility_log(output_path):
    """Generate reproducibility log markdown. (Part 5)"""
    pass  # Will be implemented next


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    print("=" * 60)
    print("MetaJudge v0.5.5.1 — Statistical Analysis Report Generator")
    print("=" * 60)

    FIGURES_DIR.mkdir(parents=True, exist_ok=True)

    # Step 1: Load data
    print("\n[1/6] Loading data...")
    data = load_all_data()
    n_cal = len(set(r["item_id"] for r in data["cal_rows"]))
    n_fb = len(set(r["item_id"] for r in data["fb_rows"]))
    print(f"  Calibration: {n_cal} items, {len(data['cal_by_model'])} models")
    print(f"  Family B: {n_fb} items, {len(data['fb_by_model'])} models")
    print(f"  Flagged items: {len(data['flagged_items'])}")

    # Step 2: Run statistics (full set)
    print("\n[2/6] Running statistical tests (full set)...")
    cal_stats = run_calibration_stats(data["cal_by_model"])
    print(f"  Calibration pairwise: {len(cal_stats['pairwise'])} pairs, "
          f"{cal_stats['n_items']} items each")

    fb_stats = run_family_b_stats(data["fb_by_model"])
    print(f"  Family B pairwise: {len(fb_stats['pairwise'])} pairs, "
          f"{fb_stats['n_items']} items each")

    # Step 2b: Sensitivity analysis (clean subset)
    print("  Running sensitivity analysis (clean subset)...")
    cal_stats_clean = run_calibration_stats(data["cal_by_model"],
                                            item_filter=data["flagged_items"])
    fb_stats_clean = run_family_b_stats(data["fb_by_model"],
                                        item_filter=data["flagged_items"])
    print(f"  Clean cal: {cal_stats_clean['n_items']} items, "
          f"Clean FB: {fb_stats_clean['n_items']} items")

    # Step 3: Per-model metrics
    print("\n[3/6] Computing per-model metrics...")
    cal_metrics = compute_per_model_calibration_metrics(data["cal_by_model"])
    fb_metrics = compute_per_model_fb_metrics(data["fb_by_model"])
    mech_stats = compute_per_mechanism_stats(data["cal_by_model"])
    bridge_corr = compute_bridge_correlations(data["cal_by_model"])
    print(f"  Mechanisms: {len(mech_stats)}")
    print(f"  Bridge correlations: {len(bridge_corr)} models")

    # Step 4: Generate visualizations
    print("\n[4/6] Generating visualizations...")
    generate_visualizations(data, cal_stats, fb_stats, cal_metrics,
                            fb_metrics, mech_stats, bridge_corr,
                            cal_stats_clean, fb_stats_clean)

    # Step 5: Generate reports
    print("\n[5/6] Generating reports...")
    generate_backgrounder(OUTPUT_DIR / "metajudge_v0551_stats_backgrounder.md")
    generate_docx_report(OUTPUT_DIR / "metajudge_v0551_stats_report.docx",
                         data, cal_stats, fb_stats, cal_metrics, fb_metrics,
                         mech_stats, bridge_corr, cal_stats_clean, fb_stats_clean)

    # Step 6: Reproducibility log
    print("\n[6/6] Generating reproducibility log...")
    generate_reproducibility_log(OUTPUT_DIR / "metajudge_v0551_stats_reproducibility.md")

    print("\n" + "=" * 60)
    print("Done. Outputs in:", OUTPUT_DIR)
    print("=" * 60)


if __name__ == "__main__":
    main()

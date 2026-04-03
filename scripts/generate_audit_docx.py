#!/usr/bin/env python3
"""Generate a per-model .docx audit report from MetaJudge v6.1 notebook outputs.

Usage:
    python scripts/generate_audit_docx.py \
      --task calibration \
      --audit-csv /path/to/calibration_item_audit_flash2.5_v6.1.csv \
      --responses-json /path/to/calibration_full_responses_flash2.5_v6.1.json \
      --items-json /path/to/metajudge_benchmark_v1.json \
      --registry-json /path/to/adjudication_registry.json \
      --output MetaJudge_Calibration_Flash2.5.docx \
      [--run-json /path/to/run.json] \
      [--justifications-md /path/to/metajudge_v51_gold_answer_justifications.md]

Supports all 4 task types: calibration, abstention, sc_c1, sc_c2.
"""

import argparse
import csv
import json
import re
import sys
from datetime import datetime
from pathlib import Path


# ── CLI ───────────────────────────────────────────────────────────────────

def parse_args():
    p = argparse.ArgumentParser(description="Generate MetaJudge v6.1 audit .docx")
    p.add_argument("--task", required=True,
                   choices=["calibration", "abstention", "sc_c1", "sc_c2"],
                   help="Task type")
    p.add_argument("--audit-csv", required=True, help="Item-level audit CSV")
    p.add_argument("--responses-json", required=True, help="Full responses JSON")
    p.add_argument("--items-json", required=True, help="Benchmark items JSON")
    p.add_argument("--registry-json", required=True, help="Adjudication registry JSON")
    p.add_argument("--output", required=True, help="Output .docx path")
    p.add_argument("--run-json", default=None, help="SDK run.json (optional, for runtime/cost)")
    p.add_argument("--justifications-md", default=None,
                   help="Gold answer justifications markdown (optional)")
    return p.parse_args()


# ── Data loading ──────────────────────────────────────────────────────────

def load_audit_csv(path):
    """Load audit CSV into list of dicts."""
    with open(path, newline="") as f:
        return list(csv.DictReader(f))


def load_json(path):
    """Load a JSON file."""
    with open(path) as f:
        return json.load(f)


def load_justifications(path):
    """Parse justifications markdown into dict keyed by item_id."""
    if path is None:
        return {}
    with open(path) as f:
        content = f.read()
    pattern = r"#### (\S+)\n.*?\*\*Justification:\*\* (.+?)(?=\n\n#### |\n\n### |\n\n## |\n\n---|\Z)"
    matches = re.findall(pattern, content, re.DOTALL)
    return {item_id: text.strip() for item_id, text in matches}


def load_run_metadata(path):
    """Extract model name, timestamps, and token/cost data from run.json."""
    if path is None:
        return None
    data = load_json(path)
    meta = {
        "model": "unknown",
        "start_time": None,
        "end_time": None,
        "input_tokens": 0,
        "output_tokens": 0,
        "input_cost": 0.0,
        "output_cost": 0.0,
    }
    # Model name
    mv = data.get("modelVersion", {})
    meta["model"] = mv.get("slug", data.get("model", "unknown"))
    # Timestamps
    meta["start_time"] = data.get("startTime")
    meta["end_time"] = data.get("endTime")
    # Token/cost aggregation across conversations
    for conv in data.get("conversations", []):
        for key in ("input_tokens", "output_tokens"):
            meta[key] += conv.get(key, 0)
        for key in ("input_tokens_cost", "output_tokens_cost"):
            target = "input_cost" if "input" in key else "output_cost"
            meta[target] += conv.get(key, 0.0)
    return meta


# ── Data merging ──────────────────────────────────────────────────────────

def build_items_lookup(items_json, task):
    """Build item_id → item dict, filtering by subfamily for C1/C2."""
    items = items_json if isinstance(items_json, list) else []
    if task == "sc_c1":
        items = [it for it in items if it.get("subfamily") == "C1"]
    elif task == "sc_c2":
        items = [it for it in items if it.get("subfamily") == "C2"]
    return {it["item_id"]: it for it in items}


def build_registry_lookup(registry_json):
    """Build item_id → registry entry dict."""
    if isinstance(registry_json, list):
        return {entry["item_id"]: entry for entry in registry_json}
    elif isinstance(registry_json, dict):
        return registry_json
    return {}


def extract_responses_lookup(responses_json, task):
    """Build item_id → full response dict from responses JSON.

    Handles both single-run ({"responses": [...]}) and
    dual-run ({"run_1": [...], "run_2": [...]}) formats.
    Returns (run1_lookup, run2_lookup_or_None).
    """
    # Strip metadata wrapper if present
    if "metadata" in responses_json:
        data = {k: v for k, v in responses_json.items() if k != "metadata"}
    else:
        data = responses_json

    run1 = data.get("run_1", data.get("responses", []))
    run2 = data.get("run_2", None)

    r1_lookup = {r["item_id"]: r for r in run1} if run1 else {}
    r2_lookup = {r["item_id"]: r for r in run2} if run2 else None

    return r1_lookup, r2_lookup


def merge_item_data(audit_rows, items_lookup, registry_lookup,
                    responses_r1, justifications):
    """Merge all data sources into a list of enriched item dicts.

    Iterates over the full items list so that items missing from the audit
    CSV (evaluation failures) still appear in the report.
    """
    audit_by_id = {row["item_id"]: row for row in audit_rows}
    merged = []

    # Include all items from the items JSON (not just audit CSV)
    all_ids = set(items_lookup.keys()) | set(audit_by_id.keys())

    for iid in sorted(all_ids):
        row = audit_by_id.get(iid, {})
        item = items_lookup.get(iid, {})
        reg = registry_lookup.get(iid, {})
        resp = responses_r1.get(iid, {})

        entry = {**row}  # Start with audit CSV fields (empty dict if failed)
        entry["item_id"] = iid
        entry["eval_failed"] = iid not in audit_by_id
        # Add item metadata
        entry["question"] = item.get("question", item.get("turn1_prompt", ""))
        entry["gold_answer_full"] = item.get("gold_answer", row.get("gold_answer", ""))
        entry["category"] = item.get("category", item.get("mechanism_primary", ""))
        entry["mechanism"] = item.get("mechanism", item.get("mechanism_primary", ""))
        entry["difficulty"] = item.get("difficulty", "")
        entry["grading_rule"] = item.get("grading_rule", reg.get("rule", ""))
        entry["subfamily"] = item.get("subfamily", "")
        entry["normative_t2_action"] = item.get("normative_t2_action", "")
        entry["stratum"] = item.get("stratum", "")
        entry["evidence_snippet"] = item.get("evidence_snippet", "")
        # Registry accepted forms
        entry["accepted_forms"] = reg.get("aliases", reg.get("accepted_forms", []))
        # Justification
        entry["justification"] = justifications.get(iid, "")
        # Full response text (prefer over truncated CSV)
        if resp:
            entry["full_model_answer"] = resp.get("model_answer",
                                                   resp.get("t1_answer", ""))
            entry["full_t1_answer"] = resp.get("t1_answer", "")
            entry["full_t2_answer"] = resp.get("t2_answer", "")
        # Gold action (abstention)
        if "gold_action" not in entry:
            entry["gold_action"] = item.get("gold_action", "")
        # Acceptable actions (abstention)
        entry["acceptable_actions"] = item.get("acceptable_actions", [])
        entry["is_false_presupposition"] = item.get("is_false_presupposition", False)

        merged.append(entry)

    return merged


# ── Report metadata ──────────────────────────────────────────────────────

TASK_DISPLAY = {
    "calibration": "Confidence Calibration",
    "abstention": "Selective Abstention",
    "sc_c1": "Intrinsic Self-Correction (C1)",
    "sc_c2": "Evidence-Assisted Self-Correction (C2)",
}


def infer_model_name(run_meta, responses_json):
    """Get model name from run.json or responses JSON metadata."""
    if run_meta and run_meta["model"] != "unknown":
        return run_meta["model"]
    if isinstance(responses_json, dict) and "metadata" in responses_json:
        return responses_json["metadata"].get("model", "unknown")
    return "unknown"


# ── Docx helpers ──────────────────────────────────────────────────────────

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# Colors
HEADER_BG = "1F4E79"
HEADER_FG = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BLUE_BG = "D6E4F0"
LIGHT_GRAY_BG = "F2F2F2"
GREEN_BG = "C6EFCE"
RED_BG = "FFC7CE"
GREEN_TEXT = RGBColor(0x00, 0x6B, 0x3C)
RED_TEXT = RGBColor(0xCC, 0x00, 0x00)


def _set_cell_bg(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_text(cell, text, size=10, bold=False, color=None, align=None):
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def _add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table with header row and alternating shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_text(cell, h, size=10, bold=True, color=HEADER_FG)
        _set_cell_bg(cell, HEADER_BG)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            _set_cell_text(cell, val, size=10)
            if r_idx % 2 == 1:
                _set_cell_bg(cell, LIGHT_GRAY_BG)

    return table


def _add_kv_table(doc, pairs):
    """Add a 2-column metric/value table."""
    return _add_table(doc, ["Metric", "Value"],
                      [[k, v] for k, v in pairs])


def _safe_float(val, default=0.0):
    """Safely parse a float from a string."""
    try:
        return float(val)
    except (ValueError, TypeError):
        return default


def _safe_bool(val):
    """Parse a boolean from string/bool."""
    if isinstance(val, bool):
        return val
    return str(val).strip().lower() in ("true", "1", "yes")


# ── Summary page builders ────────────────────────────────────────────────

def build_header_block(doc, task, model_name, merged, run_meta):
    """Add title and header metadata."""
    title = doc.add_heading(
        f"MetaJudge v6.1 — {TASK_DISPLAY[task]} Audit Report", level=1)
    for run in title.runs:
        run.font.name = "Arial"
        run.font.size = Pt(16)

    n_scored = sum(1 for m in merged if not m.get("eval_failed"))
    n_total = len(merged)
    n_failed = n_total - n_scored
    items_line = f"Items: {n_scored}/{n_total} scored"
    if n_failed:
        items_line += f" ({n_failed} evaluation failure{'s' if n_failed > 1 else ''})"

    meta_lines = [
        f"Model: {model_name}",
        f"Date: {run_meta['end_time'][:10] if run_meta and run_meta.get('end_time') else datetime.utcnow().strftime('%Y-%m-%d')}",
        f"Benchmark version: v6.1 | Grading engine: grading_v2",
        items_line,
    ]
    for line in meta_lines:
        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(11)
    doc.add_paragraph("")


def build_calibration_summary(doc, merged):
    """Performance summary for calibration task."""
    doc.add_heading("Performance Summary", level=2)
    merged = [m for m in merged if not m.get("eval_failed")]
    n = len(merged)
    correct = sum(1 for m in merged if _safe_bool(m.get("is_correct")))
    brier_scores = [_safe_float(m.get("brier_score")) for m in merged]
    mean_brier = sum(brier_scores) / n if n else 0
    confs = [_safe_float(m.get("confidence")) for m in merged]
    mean_conf = sum(confs) / n if n else 0
    overconf_wrong = sum(1 for m in merged
                         if not _safe_bool(m.get("is_correct"))
                         and _safe_float(m.get("confidence")) >= 0.9)

    FLOOR, CEIL = 0.75, 1.00
    normalized = max(0.0, min(1.0, (mean_brier - FLOOR) / (CEIL - FLOOR)))

    _add_kv_table(doc, [
        ("Accuracy", f"{correct}/{n} ({correct/n:.1%})"),
        ("Mean 1-Brier", f"{mean_brier:.4f}"),
        ("Normalized score", f"{normalized:.3f}"),
        ("Mean confidence", f"{mean_conf:.3f}"),
        ("Overconfident wrong (conf ≥ 0.9, incorrect)", str(overconf_wrong)),
    ])
    doc.add_paragraph("")


def build_abstention_summary(doc, merged):
    """Performance summary for abstention task."""
    doc.add_heading("Performance Summary", level=2)
    merged = [m for m in merged if not m.get("eval_failed")]
    n = len(merged)
    utilities = [_safe_float(m.get("utility")) for m in merged]
    mean_util = sum(utilities) / n if n else 0
    uwaa = (mean_util + 1) / 2
    FLOOR, CEIL = 0.60, 1.00
    normalized = max(0.0, min(1.0, (uwaa - FLOOR) / (CEIL - FLOOR)))
    act_correct = sum(1 for m in merged
                      if m.get("model_decision") == m.get("gold_action"))
    neg_util = sum(1 for u in utilities if u < 0)

    _add_kv_table(doc, [
        ("UWAA", f"{uwaa:.4f}"),
        ("Normalized score", f"{normalized:.3f}"),
        ("Action accuracy", f"{act_correct}/{n} ({act_correct/n:.1%})"),
        ("Negative utility items", str(neg_util)),
    ])
    doc.add_paragraph("")

    # Action cross-table
    doc.add_heading("Action Distribution", level=3)
    actions = ["answer", "clarify", "verify", "abstain"]
    matrix = {(ma, ga): 0 for ma in actions for ga in actions}
    for m in merged:
        md = m.get("model_decision", "").lower()
        ga = m.get("gold_action", "").lower()
        if (md, ga) in matrix:
            matrix[(md, ga)] += 1

    headers = [""] + [f"Gold: {a}" for a in actions]
    rows = []
    for ma in actions:
        row = [f"Model: {ma}"] + [str(matrix.get((ma, ga), 0)) for ga in actions]
        rows.append(row)
    _add_table(doc, headers, rows)
    doc.add_paragraph("")


def build_sc_summary(doc, merged, task):
    """Performance summary for C1/C2 tasks."""
    doc.add_heading("Performance Summary", level=2)
    merged = [m for m in merged if not m.get("eval_failed")]
    n = len(merged)
    t1c = sum(1 for m in merged if _safe_bool(m.get("t1_correct")))
    t2c = sum(1 for m in merged if _safe_bool(m.get("t2_correct")))
    delta = (t2c - t1c) / n if n else 0

    if task == "sc_c1":
        FLOOR, CEIL = -0.10, 0.15
    else:
        FLOOR, CEIL = -0.20, 0.20
    normalized = max(0.0, min(1.0, (delta - FLOOR) / (CEIL - FLOOR)))

    damage = sum(1 for m in merged
                 if _safe_bool(m.get("t1_correct"))
                 and not _safe_bool(m.get("t2_correct")))
    corrections = sum(1 for m in merged
                      if not _safe_bool(m.get("t1_correct"))
                      and _safe_bool(m.get("t2_correct")))

    _add_kv_table(doc, [
        ("T1 accuracy", f"{t1c}/{n} ({t1c/n:.1%})"),
        ("T2 accuracy", f"{t2c}/{n} ({t2c/n:.1%})"),
        ("T2-T1 delta", f"{delta:+.4f}"),
        ("Normalized score", f"{normalized:.3f}"),
        ("Damage events", str(damage)),
        ("Correction gains", str(corrections)),
    ])
    doc.add_paragraph("")

    # Transition summary
    doc.add_heading("Transition Summary", level=3)
    from collections import Counter
    transitions = Counter(m.get("transition", "unknown") for m in merged)
    trans_items = {}
    for m in merged:
        t = m.get("transition", "unknown")
        trans_items.setdefault(t, []).append(m["item_id"])

    headers = ["Transition", "Count", "Items"]
    rows = []
    for trans in ["maintain_correct", "correction_gain", "neutral_revision",
                  "damage", "stubborn_wrong", "failed_revision"]:
        count = transitions.get(trans, 0)
        if count == 0:
            items_str = "—"
        else:
            ids = trans_items.get(trans, [])
            if len(ids) <= 5:
                items_str = ", ".join(ids)
            else:
                items_str = ", ".join(ids[:5]) + f" and {len(ids)-5} more"
        rows.append([trans, str(count), items_str])
    _add_table(doc, headers, rows)
    doc.add_paragraph("")


def build_stochasticity_block(doc, merged, responses_r2, items_lookup,
                              registry_lookup, justifications, task):
    """Stochasticity comparison block (B, C1, C2 only)."""
    if not responses_r2:
        return

    doc.add_heading("Stochasticity (Dual-Run Comparison)", level=2)

    # Build Run 2 merged data
    audit_r2 = []
    for iid, resp in responses_r2.items():
        audit_r2.append(resp)

    if task == "abstention":
        # Run 1 metrics
        utils_1 = [_safe_float(m.get("utility")) for m in merged]
        uwaa_1 = (sum(utils_1) / len(utils_1) + 1) / 2 if utils_1 else 0.5
        FLOOR, CEIL = 0.60, 1.00
        norm_1 = max(0.0, min(1.0, (uwaa_1 - FLOOR) / (CEIL - FLOOR)))

        # Run 2 metrics — compute utility from responses
        # (simplified: just report if we have the data)
        utils_2 = [_safe_float(r.get("utility", 0)) for r in audit_r2]
        uwaa_2 = (sum(utils_2) / len(utils_2) + 1) / 2 if utils_2 else 0.5
        norm_2 = max(0.0, min(1.0, (uwaa_2 - FLOOR) / (CEIL - FLOOR)))

        _add_table(doc, ["Metric", "Run 1", "Run 2"], [
            ["UWAA", f"{uwaa_1:.4f}", f"{uwaa_2:.4f}"],
            ["Normalized score", f"{norm_1:.3f}", f"{norm_2:.3f}"],
        ])

        # Action stability
        r1_by_id = {m["item_id"]: m for m in merged}
        diffs = []
        stable = 0
        for iid, r2 in responses_r2.items():
            r1 = r1_by_id.get(iid)
            if r1:
                if r1.get("model_decision") == r2.get("model_decision"):
                    stable += 1
                else:
                    diffs.append((iid,
                                  r1.get("model_decision", "?"),
                                  r2.get("model_decision", "?")))
        total = len(responses_r2)
        doc.add_paragraph("")
        _add_kv_table(doc, [
            ("Items with identical actions", f"{stable}/{total} ({stable/total:.0%})"),
            ("Score range across runs", f"{min(norm_1, norm_2):.2f} – {max(norm_1, norm_2):.2f}"),
        ])

    elif task in ("sc_c1", "sc_c2"):
        n = len(merged)
        t1c_1 = sum(1 for m in merged if _safe_bool(m.get("t1_correct")))
        t2c_1 = sum(1 for m in merged if _safe_bool(m.get("t2_correct")))
        delta_1 = (t2c_1 - t1c_1) / n if n else 0
        dmg_1 = sum(1 for m in merged
                    if _safe_bool(m.get("t1_correct"))
                    and not _safe_bool(m.get("t2_correct")))

        FLOOR = -0.10 if task == "sc_c1" else -0.20
        CEIL = 0.15 if task == "sc_c1" else 0.20
        norm_1 = max(0.0, min(1.0, (delta_1 - FLOOR) / (CEIL - FLOOR)))

        # Run 2
        t1c_2 = sum(1 for r in audit_r2 if _safe_bool(r.get("t1_correct")))
        t2c_2 = sum(1 for r in audit_r2 if _safe_bool(r.get("t2_correct")))
        n2 = len(audit_r2)
        delta_2 = (t2c_2 - t1c_2) / n2 if n2 else 0
        norm_2 = max(0.0, min(1.0, (delta_2 - FLOOR) / (CEIL - FLOOR)))
        dmg_2 = sum(1 for r in audit_r2
                    if _safe_bool(r.get("t1_correct"))
                    and not _safe_bool(r.get("t2_correct")))

        _add_table(doc, ["Metric", "Run 1", "Run 2"], [
            ["Delta", f"{delta_1:+.4f}", f"{delta_2:+.4f}"],
            ["Normalized score", f"{norm_1:.3f}", f"{norm_2:.3f}"],
            ["Damage events", str(dmg_1), str(dmg_2)],
        ])

        # Transition stability
        r1_by_id = {m["item_id"]: m for m in merged}
        diffs = []
        stable = 0
        for iid, r2 in responses_r2.items():
            r1 = r1_by_id.get(iid)
            if r1:
                if r1.get("transition") == r2.get("transition"):
                    stable += 1
                else:
                    diffs.append((iid,
                                  r1.get("transition", "?"),
                                  r2.get("transition", "?")))
        total = len(responses_r2)
        doc.add_paragraph("")
        _add_kv_table(doc, [
            ("Items with identical transitions", f"{stable}/{total} ({stable/total:.0%})"),
            ("Score range across runs", f"{min(norm_1, norm_2):.2f} – {max(norm_1, norm_2):.2f}"),
            ("Interpretation", f"{len(diffs)} items changed between independent runs"),
        ])

        if diffs:
            doc.add_paragraph("")
            doc.add_heading("Changed Items", level=3)
            _add_table(doc, ["Item", "Run 1 transition", "Run 2 transition"],
                       [[iid, t1, t2] for iid, t1, t2 in diffs])

    doc.add_paragraph("")


def build_runtime_block(doc, run_meta):
    """Runtime and cost block from run.json."""
    if not run_meta:
        return

    doc.add_heading("Runtime and Cost", level=2)

    # Wall clock
    wall = "N/A"
    if run_meta.get("start_time") and run_meta.get("end_time"):
        try:
            fmt = "%Y-%m-%dT%H:%M:%S"
            start = run_meta["start_time"][:19]
            end = run_meta["end_time"][:19]
            t0 = datetime.strptime(start, fmt)
            t1 = datetime.strptime(end, fmt)
            secs = (t1 - t0).total_seconds()
            wall = f"{int(secs // 60)}m {int(secs % 60)}s"
        except (ValueError, TypeError):
            pass

    in_tok = run_meta.get("input_tokens", 0)
    out_tok = run_meta.get("output_tokens", 0)
    in_cost = run_meta.get("input_cost", 0)
    out_cost = run_meta.get("output_cost", 0)
    total_cost = in_cost + out_cost

    _add_kv_table(doc, [
        ("Wall clock", wall),
        ("Input tokens", f"{in_tok:,}" if in_tok else "N/A"),
        ("Output tokens", f"{out_tok:,}" if out_tok else "N/A"),
        ("Input cost", f"${in_cost:.4f}" if in_cost else "N/A"),
        ("Output cost", f"${out_cost:.4f}" if out_cost else "N/A"),
        ("Total cost", f"${total_cost:.4f}" if total_cost else "N/A"),
    ])
    doc.add_paragraph("")


# ── Item-by-item detail ───────────────────────────────────────────────────

def _add_separator(doc):
    """Add a thin horizontal rule between items."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("─" * 70)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


def _add_field(doc, label, value, bold_value=False, color=None):
    """Add a labeled field line: 'Label: value'."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run_label = p.add_run(f"{label}: ")
    run_label.font.name = "Arial"
    run_label.font.size = Pt(11)
    run_label.font.bold = True
    run_val = p.add_run(str(value))
    run_val.font.name = "Arial"
    run_val.font.size = Pt(11)
    run_val.font.bold = bold_value
    if color:
        run_val.font.color.rgb = color


def _add_block_text(doc, label, text):
    """Add a labeled block of text (question, justification, etc.)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run_label = p.add_run(f"{label}:\n")
    run_label.font.name = "Arial"
    run_label.font.size = Pt(11)
    run_label.font.bold = True
    run_text = p.add_run(str(text))
    run_text.font.name = "Arial"
    run_text.font.size = Pt(10)


def _get_answer_text(item, field, csv_field=None):
    """Get full answer text, preferring full response over truncated CSV."""
    full_key = f"full_{field}" if f"full_{field}" in item else None
    if full_key and item.get(full_key):
        return str(item[full_key])
    if csv_field and item.get(csv_field):
        return str(item[csv_field])
    return str(item.get(field, item.get(csv_field, "")))


def _correct_color(is_correct):
    """Return green for correct, red for incorrect."""
    return GREEN_TEXT if is_correct else RED_TEXT


def _add_failed_item(doc, item, task):
    """Render a failed item (present in items JSON, absent from audit CSV)."""
    _add_separator(doc)
    doc.add_heading(item["item_id"], level=3)

    _add_block_text(doc, "Question", item.get("question", ""))
    _add_field(doc, "Gold Answer", item.get("gold_answer_full", ""))

    justification = item.get("justification", "")
    _add_block_text(doc, "Justification",
                    justification if justification else "[pending]")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Model Answer: [EVALUATION FAILED — no response captured]")
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RED_TEXT

    _add_field(doc, "Correct", "—")
    _add_field(doc, "Score", "—")


def _add_calibration_item(doc, item):
    """Render one calibration item."""
    if item.get("eval_failed"):
        return _add_failed_item(doc, item, "calibration")

    _add_separator(doc)
    iid = item["item_id"]

    # Heading
    doc.add_heading(iid, level=3)

    # Metadata line
    parts = []
    if item.get("category"):
        parts.append(f"Category: {item['category']}")
    if item.get("mechanism"):
        parts.append(f"Mechanism: {item['mechanism']}")
    if item.get("grading_rule"):
        parts.append(f"Grading: {item['grading_rule']}")
    if item.get("difficulty"):
        parts.append(f"Difficulty: {item['difficulty']}")
    if parts:
        p = doc.add_paragraph(" | ".join(parts))
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Question
    _add_block_text(doc, "Question", item.get("question", ""))

    # Gold answer + accepted forms
    _add_field(doc, "Gold Answer", item.get("gold_answer_full", ""))
    accepted = item.get("accepted_forms", [])
    if accepted:
        _add_field(doc, "Accepted forms", ", ".join(str(a) for a in accepted))
    else:
        _add_field(doc, "Accepted forms", "[gold answer only]")

    # Justification
    justification = item.get("justification", "")
    _add_block_text(doc, "Justification",
                    justification if justification else "[pending]")

    # Model response
    is_correct = _safe_bool(item.get("is_correct"))
    model_answer = _get_answer_text(item, "model_answer", "model_answer")
    conf = _safe_float(item.get("confidence"))
    brier = _safe_float(item.get("brier_score"))

    _add_field(doc, "Model Answer", model_answer)
    _add_field(doc, "Confidence", f"{conf:.2f}")
    _add_field(doc, "Correct", "YES" if is_correct else "NO",
               bold_value=True, color=_correct_color(is_correct))
    _add_field(doc, "1-Brier", f"{brier:.4f}")


def _add_abstention_item(doc, item):
    """Render one abstention item."""
    if item.get("eval_failed"):
        return _add_failed_item(doc, item, "abstention")

    _add_separator(doc)
    iid = item["item_id"]

    doc.add_heading(iid, level=3)

    # Metadata
    parts = []
    if item.get("category"):
        parts.append(f"Category: {item['category']}")
    if item.get("gold_action"):
        parts.append(f"Gold action: {item['gold_action']}")
    if item.get("is_false_presupposition"):
        parts.append("FALSE PRESUPPOSITION")
    if parts:
        p = doc.add_paragraph(" | ".join(parts))
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Question
    _add_block_text(doc, "Question", item.get("question", ""))

    # Gold answer
    gold = item.get("gold_answer_full", "")
    if gold and gold.lower() not in ("n/a", ""):
        _add_field(doc, "Gold Answer", gold)

    # Justification
    justification = item.get("justification", "")
    _add_block_text(doc, "Justification",
                    justification if justification else "[pending]")

    # Model response
    decision = item.get("model_decision", "")
    model_answer = _get_answer_text(item, "model_answer", "model_answer")
    is_correct = _safe_bool(item.get("is_correct"))
    utility = _safe_float(item.get("utility"))

    _add_field(doc, "Model Decision", decision)
    if model_answer and decision == "answer":
        _add_field(doc, "Model Answer", model_answer)
    elif model_answer and decision != "answer":
        # Show explanation text for non-answer decisions
        _add_field(doc, "Model Response", model_answer)
    _add_field(doc, "Correct", "YES" if is_correct else "NO",
               bold_value=True, color=_correct_color(is_correct))
    util_color = GREEN_TEXT if utility > 0 else (RED_TEXT if utility < 0 else None)
    _add_field(doc, "Utility", f"{utility:+.1f}", color=util_color)


TRANSITION_STYLE = {
    "correction_gain":   {"label": "CORRECTION_GAIN",   "color": GREEN_TEXT, "bold": True},
    "maintain_correct":  {"label": "MAINTAIN_CORRECT",  "color": None,       "bold": False},
    "neutral_revision":  {"label": "NEUTRAL_REVISION",  "color": None,       "bold": False},
    "damage":            {"label": "DAMAGE",            "color": RED_TEXT,   "bold": True},
    "stubborn_wrong":    {"label": "STUBBORN_WRONG",    "color": RED_TEXT,   "bold": False},
    "failed_revision":   {"label": "FAILED_REVISION",   "color": RED_TEXT,   "bold": False},
}


def _add_sc_item(doc, item, task):
    """Render one C1/C2 self-correction item."""
    if item.get("eval_failed"):
        return _add_failed_item(doc, item, task)

    _add_separator(doc)
    iid = item["item_id"]
    subfamily = item.get("subfamily", task.upper().replace("SC_", ""))

    doc.add_heading(f"{iid} ({subfamily})", level=3)

    # Metadata
    parts = []
    if item.get("stratum"):
        parts.append(f"Stratum: {item['stratum']}")
    elif item.get("category"):
        parts.append(f"Category: {item['category']}")
    if item.get("normative_t2_action"):
        parts.append(f"Normative T2: {item['normative_t2_action']}")
    if item.get("difficulty"):
        parts.append(f"Difficulty: {item['difficulty']}")
    if parts:
        p = doc.add_paragraph(" | ".join(parts))
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Question
    _add_block_text(doc, "Question", item.get("question", ""))

    # Gold answer
    _add_field(doc, "Gold Answer", item.get("gold_answer_full", ""))

    # Evidence snippet (C2 only)
    evidence = item.get("evidence_snippet", "")
    if evidence and evidence.lower() not in ("none", ""):
        _add_block_text(doc, "Evidence Snippet", evidence)

    # Justification
    justification = item.get("justification", "")
    _add_block_text(doc, "Justification",
                    justification if justification else "[pending]")

    # Turn 1
    t1_correct = _safe_bool(item.get("t1_correct"))
    t1_answer = _get_answer_text(item, "t1_answer", "t1_answer")
    t1_conf = item.get("t1_confidence")
    t1_conf_str = f"{_safe_float(t1_conf):.2f}" if t1_conf not in (None, "", "None") else "—"

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Turn 1:")
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.bold = True

    _add_field(doc, "  Answer", t1_answer)
    _add_field(doc, "  Confidence", t1_conf_str)
    _add_field(doc, "  Correct", "YES" if t1_correct else "NO",
               bold_value=True, color=_correct_color(t1_correct))

    # Turn 2
    t2_correct = _safe_bool(item.get("t2_correct"))
    t2_answer = _get_answer_text(item, "t2_answer", "t2_answer")
    t2_conf = item.get("t2_confidence")
    t2_conf_str = f"{_safe_float(t2_conf):.2f}" if t2_conf not in (None, "", "None") else "—"

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Turn 2:")
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.bold = True

    _add_field(doc, "  Answer", t2_answer)
    _add_field(doc, "  Confidence", t2_conf_str)
    _add_field(doc, "  Correct", "YES" if t2_correct else "NO",
               bold_value=True, color=_correct_color(t2_correct))

    # Transition
    transition = item.get("transition", "unknown")
    style = TRANSITION_STYLE.get(transition, {"label": transition.upper(), "color": None, "bold": False})
    _add_field(doc, "Transition", style["label"],
               bold_value=style["bold"], color=style["color"])

    # Edit similarity
    sim = item.get("t1_t2_similarity")
    if sim not in (None, "", "None"):
        _add_field(doc, "T1→T2 similarity", f"{_safe_float(sim):.3f}")


# ── Document assembly ────────────────────────────────────────────────────

def create_document(task, model_name, merged, responses_r2, items_lookup,
                    registry_lookup, justifications, run_meta):
    """Build the full docx document."""
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # Summary page
    build_header_block(doc, task, model_name, merged, run_meta)

    if task == "calibration":
        build_calibration_summary(doc, merged)
    elif task == "abstention":
        build_abstention_summary(doc, merged)
    elif task in ("sc_c1", "sc_c2"):
        build_sc_summary(doc, merged, task)

    # Stochasticity (B, C1, C2 only)
    if task != "calibration":
        build_stochasticity_block(doc, merged, responses_r2, items_lookup,
                                  registry_lookup, justifications, task)

    build_runtime_block(doc, run_meta)

    # Page break before item details
    doc.add_page_break()

    doc.add_heading("Item-by-Item Detail", level=1)
    for item in merged:
        if task == "calibration":
            _add_calibration_item(doc, item)
        elif task == "abstention":
            _add_abstention_item(doc, item)
        elif task in ("sc_c1", "sc_c2"):
            _add_sc_item(doc, item, task)

    return doc


# ── Main ──────────────────────────────────────────────────────────────────

def main():
    args = parse_args()

    # Load all data
    print(f"Loading data for task: {args.task}")
    audit_rows = load_audit_csv(args.audit_csv)
    responses_json = load_json(args.responses_json)
    items_json = load_json(args.items_json)
    registry_json = load_json(args.registry_json)
    justifications = load_justifications(args.justifications_md)
    run_meta = load_run_metadata(args.run_json)

    # Build lookups
    items_lookup = build_items_lookup(items_json, args.task)
    registry_lookup = build_registry_lookup(registry_json)
    responses_r1, responses_r2 = extract_responses_lookup(responses_json, args.task)

    # Merge
    merged = merge_item_data(audit_rows, items_lookup, registry_lookup,
                             responses_r1, justifications)

    model_name = infer_model_name(run_meta, responses_json)

    print(f"  Model: {model_name}")
    print(f"  Items: {len(merged)}")
    print(f"  Justifications matched: {sum(1 for m in merged if m['justification'])}/{len(merged)}")
    print(f"  Run 2 data: {'yes' if responses_r2 else 'no'}")

    # Generate docx
    doc = create_document(args.task, model_name, merged, responses_r2,
                          items_lookup, registry_lookup, justifications, run_meta)
    doc.save(args.output)
    print(f"\nSaved: {args.output}")


if __name__ == "__main__":
    main()
